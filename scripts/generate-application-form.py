# Generates the Research Fellow / Associate Research Fellow / Research
# Assistant Application Form as a fillable Microsoft Word document
# (replacing the old PDF version), in English, Simplified Chinese, and
# Traditional Chinese.
#
# Usage: python scripts/generate-application-form.py
# Requires: python-docx  (pip install python-docx)
# Needs the same rasterized logo PNGs as generate-document-templates.py
# (see PRI_LOGO_ASSETS_DIR / .logo-assets note in that script).

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(ROOT, 'public', 'forms')
LOGO_ASSETS_DIR = os.environ.get('PRI_LOGO_ASSETS_DIR', os.path.join(ROOT, '.logo-assets'))
LOGO_HEADER = os.path.join(LOGO_ASSETS_DIR, 'logo-header.png')
LOGO_COVER = os.path.join(LOGO_ASSETS_DIR, 'logo-cover.png')

TEXT_MAIN = RGBColor(0x1F, 0x23, 0x26)
TEXT_SECONDARY = RGBColor(0x5C, 0x63, 0x68)
TEXT_MUTED = RGBColor(0x8A, 0x91, 0x96)
ACCENT_STEEL = RGBColor(0x37, 0x5A, 0x6B)
BORDER_LIGHT = 'DADDDD'
BG_SOFT = 'F1F2F2'
BG_DARK = '1F2326'
WHITE = 'FFFFFF'

CJK_SANS = {'en': 'Calibri', 'zh-cn': 'Microsoft YaHei', 'zh-tw': 'Microsoft JhengHei'}
CJK_SERIF = {'en': 'Georgia', 'zh-cn': 'SimSun', 'zh-tw': 'PMingLiU'}
CJK_MONO = {'en': 'Consolas', 'zh-cn': 'Microsoft YaHei', 'zh-tw': 'Microsoft JhengHei'}
LANG_SUFFIX = {'en': 'EN', 'zh-cn': 'CN', 'zh-tw': 'CN-TW'}
CHECK = '☐'  # ☐


def set_font(run, font_name, size=None, bold=None, italic=None, color=None, east_asian=None):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)
    rfonts.set(qn('w:eastAsia'), east_asian or font_name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def no_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tc_pr.append(borders)


def add_bottom_rule(paragraph, color=BORDER_LIGHT, sz=6, space=4):
    if isinstance(color, RGBColor):
        color = str(color)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    set_font(run, 'Consolas', 9, color=TEXT_MUTED)
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    i = OxmlElement('w:instrText'); i.set(qn('xml:space'), 'preserve'); i.text = 'PAGE'
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(i); run._r.append(e)


class Ctx:
    def __init__(self, lang):
        self.lang = lang
        self.sans = 'Calibri'
        self.serif = 'Georgia'
        self.mono = 'Consolas'
        self.ea_sans = CJK_SANS[lang]
        self.ea_serif = CJK_SERIF[lang]
        self.ea_mono = CJK_MONO[lang]


# ---------------------------------------------------------------------------
# Layout primitives
# ---------------------------------------------------------------------------

def build_header_footer(doc, ctx, official_label, footer_line, instructions_label):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    header = section.header
    h_table = header.add_table(rows=1, cols=2, width=Cm(16.0))
    h_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left_cell, right_cell = h_table.rows[0].cells
    left_p = left_cell.paragraphs[0]
    run = left_p.add_run()
    run.add_picture(LOGO_HEADER, width=Cm(3.4))
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = right_p.add_run(official_label)
    set_font(r, ctx.mono, 9, color=TEXT_MUTED, east_asian=ctx.ea_mono)
    rule_p = header.add_paragraph()
    add_bottom_rule(rule_p, color=BORDER_LIGHT, sz=6)

    footer = section.footer
    f_table = footer.add_table(rows=1, cols=2, width=Cm(16.0))
    f_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fl_cell, fr_cell = f_table.rows[0].cells
    fl_p = fl_cell.paragraphs[0]
    r = fl_p.add_run(footer_line)
    set_font(r, ctx.mono, 8.5, color=TEXT_MUTED, east_asian=ctx.ea_mono)
    fr_p = fr_cell.paragraphs[0]
    fr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fr_p.add_run(instructions_label + ' · ')
    set_font(r, ctx.mono, 8.5, color=TEXT_MUTED, east_asian=ctx.ea_mono)
    add_page_number_field(fr_p)


def add_cover(doc, ctx, form_no, title_main, title_sub, desc, issued_by_label, issued_by_value, appno_label):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(LOGO_COVER, width=Cm(3.6))
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(form_no)
    set_font(r, ctx.mono, 10, color=TEXT_MUTED, east_asian=ctx.ea_mono)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title_main)
    set_font(r, ctx.serif, 34, bold=True, color=TEXT_MAIN, east_asian=ctx.ea_serif)
    if title_sub:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(18)
        r = p2.add_run(title_sub)
        set_font(r, ctx.serif, 15, color=TEXT_MUTED, east_asian=ctx.ea_serif)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(16)
    add_bottom_rule(rule, color=TEXT_MAIN, sz=12, space=1)
    rr = rule.add_run(' ')
    set_font(rr, ctx.sans, 1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(90)
    r = p.add_run(desc)
    set_font(r, ctx.sans, 12.5, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)

    rule2 = doc.add_paragraph()
    add_bottom_rule(rule2, color=BORDER_LIGHT, sz=6)
    rule2.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    lc, rc = table.rows[0].cells
    no_borders(lc); no_borders(rc)
    lp = lc.paragraphs[0]
    r = lp.add_run(issued_by_label)
    set_font(r, ctx.mono, 9, color=TEXT_MUTED, east_asian=ctx.ea_mono)
    lp2 = lc.add_paragraph()
    r = lp2.add_run(issued_by_value)
    set_font(r, ctx.sans, 11, color=TEXT_MAIN, east_asian=ctx.ea_sans)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = rp.add_run(appno_label)
    set_font(r, ctx.mono, 9, color=TEXT_MUTED, east_asian=ctx.ea_mono)


def part_heading(doc, ctx, part_label, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(part_label)
    set_font(r, ctx.mono, 10, color=TEXT_MUTED, east_asian=ctx.ea_mono)
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run(title)
    set_font(r, ctx.serif, 19, bold=True, color=TEXT_MAIN, east_asian=ctx.ea_serif)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(14)
    add_bottom_rule(rule, color=TEXT_MAIN, sz=10, space=1)


def sub_heading(doc, ctx, text, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, ctx.sans, 12, bold=True, color=TEXT_MAIN, east_asian=ctx.ea_sans)
    if note:
        np = doc.add_paragraph()
        np.paragraph_format.space_after = Pt(6)
        r = np.add_run(note)
        set_font(r, ctx.sans, 9.5, italic=True, color=TEXT_MUTED, east_asian=ctx.ea_sans)


def body_text(doc, ctx, text, space_after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, ctx.sans, 10.5, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)


def note_box(doc, ctx, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, BG_SOFT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, ctx.sans, 9.5, italic=True, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _field_cell(cell, ctx, label):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_font(r, ctx.mono, 8.5, color=TEXT_MUTED, east_asian=ctx.ea_mono)
    blank = cell.add_paragraph()
    blank.paragraph_format.space_after = Pt(10)
    add_bottom_rule(blank, color=BORDER_LIGHT, sz=6, space=1)
    br = blank.add_run(' ')
    set_font(br, ctx.sans, 8)


def field_pair(doc, ctx, label_a, label_b):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    ca, cb = table.rows[0].cells
    no_borders(ca); no_borders(cb)
    _field_cell(ca, ctx, label_a)
    _field_cell(cb, ctx, label_b)


def field_single(doc, ctx, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_font(r, ctx.mono, 8.5, color=TEXT_MUTED, east_asian=ctx.ea_mono)
    blank = doc.add_paragraph()
    blank.paragraph_format.space_after = Pt(10)
    add_bottom_rule(blank, color=BORDER_LIGHT, sz=6, space=1)
    br = blank.add_run(' ')
    set_font(br, ctx.sans, 8)


def checkbox_grid(doc, ctx, options, cols=3):
    rows = (len(options) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = True
    idx = 0
    for ri in range(rows):
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            no_borders(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(6)
            if idx < len(options):
                r = p.add_run(CHECK + '  ' + options[idx])
                set_font(r, ctx.sans, 10, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)
            idx += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def checkbox_inline(doc, ctx, prompt, options):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(prompt)
    set_font(r, ctx.sans, 10.5, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    for opt in options:
        r = p2.add_run(CHECK + '  ' + opt + '     ')
        set_font(r, ctx.sans, 10, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)


def lines(doc, ctx, count=2):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        add_bottom_rule(p, color=BORDER_LIGHT, sz=6, space=1)
        r = p.add_run(' ')
        set_font(r, ctx.sans, 8)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def numbered_lines(doc, ctx, start, count):
    for i in range(start, start + count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f'{i}. ')
        set_font(r, ctx.sans, 10.5, color=TEXT_MUTED, east_asian=ctx.ea_sans)
        add_bottom_rule(p, color=BORDER_LIGHT, sz=6, space=1)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def boxed_repeat(doc, ctx, label, body_fn):
    table = doc.add_table(rows=1, cols=1)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement('w:cantSplit'))
    cell = table.rows[0].cells[0]
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4'); el.set(qn('w:space'), '0'); el.set(qn('w:color'), BORDER_LIGHT)
        borders.append(el)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label)
    set_font(r, ctx.mono, 9, bold=True, color=TEXT_MUTED, east_asian=ctx.ea_mono)
    body_fn(cell)
    last = cell.paragraphs[-1]
    last.paragraph_format.space_after = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def dark_banner(doc, ctx, title, subtitle):
    table = doc.add_table(rows=1, cols=1)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement('w:cantSplit'))
    cell = table.rows[0].cells[0]
    shade_cell(cell, BG_DARK)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, ctx.sans, 13, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), east_asian=ctx.ea_sans)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    r = p2.add_run(subtitle)
    set_font(r, ctx.sans, 9.5, color=RGBColor(0xC5, 0xC9, 0xC9), east_asian=ctx.ea_sans)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def signature_row(doc, ctx, sig_label, date_label, y, m, d):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    ca, cb = table.rows[0].cells
    no_borders(ca); no_borders(cb)
    _field_cell(ca, ctx, sig_label)
    p = cb.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(date_label)
    set_font(r, ctx.mono, 8.5, color=TEXT_MUTED, east_asian=ctx.ea_mono)
    blank = cb.add_paragraph()
    blank.paragraph_format.space_after = Pt(10)
    r1 = blank.add_run(f'{y}          {m}          {d}')
    set_font(r1, ctx.sans, 9, color=TEXT_MUTED, east_asian=ctx.ea_sans)
    add_bottom_rule(blank, color=BORDER_LIGHT, sz=6, space=1)


# ---------------------------------------------------------------------------
# Content (EN / zh-cn / zh-tw)
# ---------------------------------------------------------------------------

T = {
    'en': dict(
        form_no='FORM NO. PRI-FRM-EN',
        title_main='Research Fellow',
        title_sub='Application Form',
        cover_desc='For applications to Research Fellow, Associate Research Fellow, Research Assistant, and related academic appointments at the Panorama Research Institute.',
        issued_by_label='ISSUED BY', issued_by_value='Panorama Research Institute · Panorama Scholarly Group',
        appno_label='APPLICATION NO. (INTERNAL)',
        official_label='Official Template', footer_line='Panorama Research Institute  ·  research.panorama-sg.com  ·  Official Template',
        instructions_label='Instructions',
        p1_label='PART ONE', p1_title='1. Application Notes',
        p1_body1="The Panorama Research Institute is a research and academic development platform of Panorama Scholarly Group. Its work encompasses academic research, the organization of research projects, scholarly publishing studies, indexing and evaluation research, policy and social research, research on artificial intelligence and future society, academic conferences, think-tank reports, and international collaboration.",
        p1_body2="This form is used to apply for an academic appointment as a Research Fellow, Associate Research Fellow, Research Assistant, or similar role. Applicants must ensure that all submitted information and materials are truthful, accurate, and complete. The Institute conducts a comprehensive evaluation based on the applicant's educational background, research experience, scholarly output, alignment of research direction, future research plans, and availability to contribute to the Institute's work.",
        p1_body3="Submission of this form does not in itself constitute an offer of employment, a labour or employment relationship, or any commitment to remuneration. Final appointment outcomes, terms, responsibilities, entitlements, modes of collaboration, and fee arrangements are governed by the Institute's official written notice, appointment documents, collaboration agreements, or project agreements.",
        p1_note="The application number is assigned internally by the Institute upon receipt; the applicant need not complete it.",
        p2_label='PART TWO', p2_title='2. Intake Information',
        f_appno='Application No.', f_appdate='Application date', f_y='Y', f_m='M', f_d='D',
        cat_label='Category applied for',
        cat_opts=['Research Fellow', 'Associate Research Fellow', 'Research Assistant', 'Young Research Fellow', 'Visiting Research Fellow', 'Part-time Research Fellow', 'Distinguished Research Fellow', 'Other'],
        term_label='Proposed term', term_opts=['1 year', '2 years', '3 years', 'Other'],
        mode_label='Proposed mode of participation', mode_opts=['Remote participation', 'Project collaboration', 'Academic appointment', 'Combined online & in-person', 'Other'],
        p3_label='PART THREE', p3_title='3. Applicant Information',
        f_fullname='Full name', f_nameLatin='Name in Latin script', f_nationality='Nationality / region', f_residence='Country / region of residence',
        f_tel='Telephone', f_email='Email', f_institution='Current institution', f_faculty='Faculty / department',
        f_position='Current position / title', f_degree='Highest degree', f_researchareas='Primary research areas',
        f_orcid='ORCID', f_scholar='Google Scholar / profile', f_website='Personal website', f_address='Mailing address',
        p4_label='PART FOUR', p4_title='4. Education', p4_note='Please list starting from your highest degree.',
        edu_label='EDUCATION', f_dates='Dates', f_institution2='Institution', f_country='Country / region', f_field='Field of study', f_degree2='Degree',
        p5_label='PART FIVE', p5_title='5. Employment & Academic Appointments',
        pos_label='POSITION', f_org='Organization', f_role='Role / title', f_mainresp='Main responsibilities or academic work:',
        p6_label='PART SIX', p6_title='6. Proposed Research Direction',
        center_label='Research area or center you wish to join',
        center_opts=['Scholarly Publishing Studies', 'Journal Indexing & Evaluation', 'Public Policy & Social Research', 'AI & Future Society', 'Education & Learning Research', 'Arts, Culture & Embodiment', 'International & Regional Studies', 'Digital & Platform Governance', 'Youth Issues & Social Development', 'Other'],
        keywords_label='Research keywords',
        p6_desc="Please briefly describe your main research direction, research questions, methods, and their relevance to the Panorama Research Institute:",
        p7_label='PART SEVEN', p7_title='7. Principal Research Output',
        p7_sub1='1. Representative publications', p7_note1='List your most representative papers from the past five years, including authors, title, journal, year, volume/issue, pages, and DOI or link.',
        p7_sub2='2. Books, textbooks, translations, or research reports',
        p7_sub3='3. Other academic output', p7_note3='May include conference papers, databases, index systems, course development, software tools, policy recommendations, media outreach, public service, or knowledge-transfer outcomes.',
        p8_label='PART EIGHT', p8_title='8. Research Projects & Academic Impact',
        proj_label='PROJECT', f_projtitle='Project title', f_funder='Funder / commissioner', f_yourrole='Your role',
        status_label='Status:', status_opts=['Ongoing', 'Completed', 'Under application', 'Other'],
        honours_label='Academic honours, awards, social impact, or academic service:',
        p9_label='PART NINE', p9_title='9. Statement of Application',
        p9_sub1='1. Reasons for applying to the Institute', p9_note1="Describe your motivation, your fit with the Institute's mission, and the academic work you hope to pursue through its platform.",
        p9_sub2='2. Research plan for the next 1–3 years', p9_note2='Describe future research themes, expected outcomes, publication plans, project plans, report plans, and how they connect with the Institute.',
        p9_sub3='3. Contributions you can offer the Institute', p9_note3='May include publications, grant applications, research reports, policy briefs, conferences, international collaboration, journal development, database building, course development, and academic outreach.',
        p10_label='PART TEN', p10_title='10. Commitment & Time Allocation',
        time_label='Estimated time available per month', time_opts=['1–3 hours', '4–8 hours', '9–15 hours', 'More than 15 hours', 'Other'],
        activities_label='Institute activities you can take part in',
        activities_opts=['Research projects', 'Research reports / policy briefs', 'Working papers', 'Edited volumes / monographs', 'Conferences / seminars', 'Lectures / training courses', 'Journal development / peer review', 'Database / index building', 'International collaboration', 'Other'],
        commit_label='Minimum annual contribution intended',
        commit_opts=['Submit at least one research output per year', 'Take part in at least one Institute event per year', 'Take part in at least one research project per year', 'To be agreed with the Institute', 'Other'],
        p11_label='PART ELEVEN', p11_title='11. Languages & Professional Skills',
        lang_label='Language ability', lang_opts=['Chinese', 'English', 'German', 'Korean', 'Japanese', 'Other'],
        skills_label='Research methods & technical skills',
        skills_opts=['Quantitative', 'Qualitative', 'Mixed methods', 'Comparative', 'Case study', 'Bibliometrics', 'Text analysis', 'Policy analysis', 'SPSS / Stata', 'R / Python', 'LaTeX', 'NVivo', 'GIS', 'Database building', 'AI tools', 'Websites / digital platforms', 'Other'],
        p12_label='PART TWELVE', p12_title='12. Ethics & Conflict of Interest',
        eth_q1='Are there any ongoing matters of academic misconduct, publication ethics, or research ethics involving you?',
        eth_q2='Is there any conflict of interest that could affect an appointment, project collaboration, or academic judgment?',
        eth_q3="Do you agree to abide by the Institute's Charter, academic standards, research ethics, publication ethics, data and documentation standards, conflict of interest policy, and name and logo usage rules?",
        yes_no=['No', 'Yes, please specify'], yes_no2=['Yes', 'No'],
        p13_label='PART THIRTEEN', p13_title='13. Checklist of Materials',
        req_label='Required', req_opts=['Research Fellow application form', 'Curriculum vitae', 'Proof of highest degree or enrolment', 'Representative papers or output', 'Research plan or statement', 'Academic profile / publication list'],
        opt_label='Optional', opt_opts=['Letters of recommendation', 'Award certificates', 'Project documentation', 'Proof of employment', 'Publishing contract / proof of publication', 'Other supporting materials'],
        p14_label='PART FOURTEEN', p14_title="14. Applicant's Declaration",
        decl_intro='I solemnly declare that:',
        decl_items=[
            'This form and all accompanying materials are submitted by me, and their contents are truthful, accurate, and complete.',
            "I understand that the Panorama Research Institute may conduct eligibility review, academic evaluation, research-alignment assessment, and appointment review based on my materials, and may require supplementary documentation.",
            "I understand that the Panorama Research Institute is a research and academic development platform of Panorama Scholarly Group and is not an independent legal entity.",
            "I understand that submitting this form or passing review does not in itself constitute a labour relationship, employment relationship, commitment to remuneration, or any other legally binding engagement.",
            "I undertake to abide by the Institute's Charter, academic standards, research ethics, publication ethics, data and documentation standards, conflict of interest policy, and name and logo usage rules.",
            "I undertake to observe the principles of academic integrity when participating in the Institute's projects, publications, conferences, reports, databases, index building, and collaborative activities, and to be free of plagiarism, fabrication, falsification, duplicate publication, improper authorship, or fraudulent peer review.",
            "I consent to the Institute's reasonable use of my submitted information and materials in the course of application review, academic evaluation, appointment management, academic-record keeping, platform display, and subsequent communication.",
            "I understand that the final appointment outcome, term, responsibilities, entitlements, and mode of collaboration are governed by the Institute's official written notice, appointment documents, collaboration agreement, or project agreement.",
        ],
        sig_label='Applicant’s signature', date_label='Date',
        internal_title='Internal Review Section', internal_subtitle='INTERNAL REVIEW ONLY · The following is completed by the Panorama Research Institute. Applicants need not complete this section.',
        p15_label='', p15_title='15. Preliminary Review',
        completeness_label='Completeness of materials', completeness_opts=['Complete', 'Largely complete', 'Needs supplement', 'Incomplete'],
        eligibility_label='Eligibility', eligibility_opts=['Eligible', 'Largely eligible', 'Needs verification', 'Ineligible'],
        background_label='Academic background', background_opts=['Excellent', 'Good', 'Average', 'Below requirement'],
        alignment_label='Research alignment', alignment_opts=['Highly aligned', 'Fairly aligned', 'Average', 'Not aligned'],
        prelim_comments_label='Preliminary comments:', prelim_supplement_label='Materials to supplement or items to revise:',
        prelim_conclusion_label='Preliminary conclusion', prelim_conclusion_opts=['Pass', 'Conditional pass', 'Deferred', 'Fail'],
        reviewer_label='Reviewer',
        p16_label='INTERNAL REVIEW · CONT.', p16_title='16. Academic Review',
        committee_label='Comments of the Academic Committee / Research Lead:',
        recommend_label='Recommended appointment',
        recommend_opts=['Appoint as Research Fellow', 'Appoint as Associate Research Fellow', 'Appoint as Research Assistant', 'Appoint as Young Research Fellow', 'Redirect to Visiting Fellow application', 'Defer', 'Not recommended'],
        recterm_label='Recommended term', recterm_opts=['1 year', '2 years', '3 years', 'Other'],
        reccenter_label='Recommended area / center', reccenter_opts=['Scholarly Publishing Studies', 'Journal Indexing & Evaluation', 'Public Policy & Social Research', 'AI & Future Society', 'Arts, Culture & Embodiment', 'To be determined'],
        finaldecision_label='Final decision', finaldecision_opts=['Approved', 'Conditionally approved', 'Deferred', 'Not approved'],
        academiclead_label='Academic lead signature', institutedirector_label='Institute director signature',
        closing_tagline='A Research and Academic Development Platform of Panorama Scholarly Group',
        closing_contact='Email: research@panorama-sg.com  ·  Website: research.panorama-sg.com',
    ),
}

T['zh-cn'] = dict(
    form_no='表格编号 PRI-FRM-CN',
    title_main='研究员申请表', title_sub='Research Fellow Application Form',
    cover_desc='适用于申请全景研究院研究员、副研究员、研究助理及相关学术任命。',
    issued_by_label='发布机构', issued_by_value='全景研究院 · 全景学术集团',
    appno_label='申请编号（内部填写）',
    official_label='官方模板', footer_line='全景研究院  ·  research.panorama-sg.com  ·  官方模板',
    instructions_label='填写说明',
    p1_label='第一部分', p1_title='一、申请说明',
    p1_body1='全景研究院是全景学术集团旗下的研究与学术发展平台，主要承担学术研究、研究项目组织、学术出版研究、索引与评价研究、政策与社会研究、人工智能与未来社会研究、学术会议、智库报告及国际合作等相关工作。',
    p1_body2='本申请表用于申请研究院研究员、副研究员、研究助理等学术任命。申请人应确保所提交信息和材料真实、准确、完整。研究院将根据申请人的教育背景、研究经历、科研成果、研究方向匹配度、未来研究计划及可参与研究院工作的情况进行综合评估。',
    p1_body3='提交本申请表并不当然构成录用、聘用、劳动关系、雇佣关系或薪酬支付承诺。具体任命结果、任期、职责、权益、合作方式及费用安排，以研究院正式书面通知、任命文件、合作协议或项目协议为准。',
    p1_note='申请编号由研究院内部受理后统一填写，申请人无需填写。',
    p2_label='第二部分', p2_title='二、受理信息',
    f_appno='申请编号', f_appdate='申请日期', f_y='年', f_m='月', f_d='日',
    cat_label='申请类别',
    cat_opts=['研究员', '副研究员', '研究助理', '青年研究员', '访问研究员', '兼职研究员', '特聘研究员', '其他'],
    term_label='拟申请任期', term_opts=['1 年', '2 年', '3 年', '其他'],
    mode_label='拟参与方式', mode_opts=['远程参与', '项目合作', '学术任命', '线上线下结合', '其他'],
    p3_label='第三部分', p3_title='三、申请人基本信息',
    f_fullname='中文姓名', f_nameLatin='英文姓名 / 拼音', f_nationality='国籍 / 地区', f_residence='现居国家 / 地区',
    f_tel='联系电话', f_email='电子邮箱', f_institution='现任单位', f_faculty='所属院系 / 部门',
    f_position='现任职务 / 职称', f_degree='最高学历 / 学位', f_researchareas='主要研究方向',
    f_orcid='ORCID', f_scholar='Google Scholar / 学术主页', f_website='个人网站 / 主页', f_address='通讯地址',
    p4_label='第四部分', p4_title='四、教育背景', p4_note='请从最高学历开始填写。',
    edu_label='教育经历', f_dates='起止时间', f_institution2='院校名称', f_country='国家 / 地区', f_field='专业 / 研究方向', f_degree2='学位',
    p5_label='第五部分', p5_title='五、工作经历与学术任职',
    pos_label='经历', f_org='单位 / 机构', f_role='职务 / 身份', f_mainresp='主要职责或学术工作：',
    p6_label='第六部分', p6_title='六、拟申请研究方向',
    center_label='拟申请加入的研究方向或研究中心',
    center_opts=['学术出版研究', '期刊索引与评价研究', '公共政策与社会研究', '人工智能与未来社会', '教育与学习研究', '艺术、文化与具身研究', '国际关系与区域研究', '数字治理与平台治理', '青年议题与社会发展', '其他'],
    keywords_label='研究关键词',
    p6_desc='请简要说明您的主要研究方向、研究问题、研究方法，以及与全景研究院研究平台的关联：',
    p7_label='第七部分', p7_title='七、主要科研成果',
    p7_sub1='1. 代表性论文', p7_note1='请列出近五年或最具代表性的论文，包括作者、题名、期刊名称、年份、卷期、页码、DOI 或链接。',
    p7_sub2='2. 著作、教材、译著或研究报告',
    p7_sub3='3. 其他学术成果', p7_note3='可填写会议论文、数据库、指数体系、课程建设、软件工具、政策建议、媒体传播、社会服务或知识转化成果。',
    p8_label='第八部分', p8_title='八、科研项目与学术影响',
    proj_label='科研项目', f_projtitle='项目名称', f_funder='资助 / 委托单位', f_yourrole='本人角色',
    status_label='项目状态：', status_opts=['在研', '已结项', '申报中', '其他'],
    honours_label='学术荣誉、奖励、社会影响或学术服务经历：',
    p9_label='第九部分', p9_title='九、申请陈述',
    p9_sub1='1. 申请加入研究院的原因', p9_note1='请说明申请动机、与研究院定位的契合度，以及希望通过研究院平台开展的学术工作。',
    p9_sub2='2. 未来 1—3 年研究计划', p9_note2='请说明未来研究主题、预期成果、论文计划、项目计划、报告计划及与研究院平台的结合方式。',
    p9_sub3='3. 可为研究院提供的学术贡献', p9_note3='可包括论文发表、课题申报、研究报告、政策简报、学术会议、国际合作、期刊建设、数据库建设、课程开发、学术传播等。',
    p10_label='第十部分', p10_title='十、参与承诺与时间安排',
    time_label='预计每月可投入时间', time_opts=['1—3 小时', '4—8 小时', '9–15 小时', '15 小时以上', '其他'],
    activities_label='可参与的研究院事务',
    activities_opts=['研究项目', '研究报告 / 政策简报', '工作论文', '论文集 / 专著', '学术会议 / 研讨会', '讲座 / 培训课程', '期刊建设 / 审稿', '数据库 / 指数建设', '国际合作', '其他'],
    commit_label='年度最低贡献意向',
    commit_opts=['每年至少提交 1 项研究成果', '每年至少参与 1 次研究院活动', '每年至少参与 1 个研究项目', '根据研究院安排协商确定', '其他'],
    p11_label='第十一部分', p11_title='十一、语言能力与专业技能',
    lang_label='语言能力', lang_opts=['中文', '英文', '德文', '韩文', '日文', '其他'],
    skills_label='研究方法与技术能力',
    skills_opts=['定量研究', '定性研究', '混合研究', '比较研究', '案例研究', '文献计量', '文本分析', '政策分析', 'SPSS / Stata', 'R / Python', 'LaTeX', 'NVivo', 'GIS', '数据库建设', 'AI 工具', '网站 / 数字平台', '其他'],
    p12_label='第十二部分', p12_title='十二、学术伦理与利益冲突声明',
    eth_q1='是否存在正在处理中的学术不端、出版伦理或研究伦理问题：',
    eth_q2='是否存在可能影响研究院任命、项目合作或学术判断的利益冲突：',
    eth_q3='是否愿意遵守研究院章程、学术规范、研究伦理、出版伦理、数据与文档规范、利益冲突政策及名称标识使用规范：',
    yes_no=['否', '是，请说明'], yes_no2=['是', '否'],
    p13_label='第十三部分', p13_title='十三、申请材料清单',
    req_label='必交材料', req_opts=['研究员申请表', '个人简历', '最高学位证明或在读证明', '代表性论文或研究成果', '研究计划或申请陈述', '学术主页 / 出版物清单'],
    opt_label='选交材料', opt_opts=['推荐信', '获奖证明', '项目证明', '任职或在职证明', '出版合同 / 发表证明', '其他补充材料'],
    p14_label='第十四部分', p14_title='十四、申请人声明',
    decl_intro='本人郑重声明：',
    decl_items=[
        '本申请表及所附材料均由本人提交，内容真实、准确、完整。',
        '本人了解全景研究院有权依据申请材料进行资格审查、学术评估、研究方向匹配及任职审核，并可要求本人补充相关证明材料。',
        '本人了解全景研究院是全景学术集团旗下研究与学术发展平台，并非独立法人实体。',
        '本人了解提交本表或通过审核，并不当然构成劳动关系、雇佣关系、薪酬支付承诺或其他法律意义上的聘用关系。',
        '本人承诺遵守研究院章程、学术规范、研究伦理、出版伦理、数据与文档规范、利益冲突政策及名称标识使用规范。',
        '本人承诺在参与研究院相关项目、出版、会议、报告、数据库、指数建设及合作活动时，遵守学术诚信原则，不存在抄袭、伪造、篡改、重复发表、不当署名、虚假同行评审等学术不端行为。',
        '本人同意研究院在申请审核、学术评估、任职管理、学术档案建立、平台展示及后续联络过程中合理使用本人提交的信息与材料。',
        '本人了解最终任命结果、任期、职责、权益及合作方式，以研究院正式书面通知、任命文件、合作协议或项目协议为准。',
    ],
    sig_label='申请人签名', date_label='日期',
    internal_title='研究院内部审核部分', internal_subtitle='INTERNAL REVIEW ONLY · 以下内容由全景研究院内部填写，申请人无需填写。',
    p15_label='', p15_title='十五、初审意见',
    completeness_label='材料完整性', completeness_opts=['完整', '基本完整', '需补充', '不完整'],
    eligibility_label='申请资格', eligibility_opts=['符合', '基本符合', '需进一步核实', '不符合'],
    background_label='学术背景', background_opts=['优秀', '良好', '一般', '不符合要求'],
    alignment_label='研究方向匹配度', alignment_opts=['高度匹配', '较匹配', '一般', '不匹配'],
    prelim_comments_label='初审意见：', prelim_supplement_label='需补充材料或修改事项：',
    prelim_conclusion_label='初审结论', prelim_conclusion_opts=['通过', '有条件通过', '暂缓', '不通过'],
    reviewer_label='初审人',
    p16_label='内部审核 · 续', p16_title='十六、学术复审意见',
    committee_label='学术委员会 / 研究负责人意见：',
    recommend_label='学术任命建议',
    recommend_opts=['建议任命为研究员', '建议任命为副研究员', '建议任命为研究助理', '建议任命为青年研究员', '建议转为访问研究员申请', '建议暂缓', '不建议任命'],
    recterm_label='建议任期', recterm_opts=['1 年', '2 年', '3 年', '其他'],
    reccenter_label='建议归属研究方向 / 研究中心', reccenter_opts=['学术出版研究', '期刊索引与评价研究', '公共政策与社会研究', '人工智能与未来社会', '艺术、文化与具身研究', '暂不指定'],
    finaldecision_label='最终审核结论', finaldecision_opts=['批准', '有条件批准', '暂缓', '不批准'],
    academiclead_label='学术负责人签名', institutedirector_label='研究院负责人签名',
    closing_tagline='全景学术集团旗下研究与学术发展平台',
    closing_contact='Email: research@panorama-sg.com  ·  Website: research.panorama-sg.com',
)

T['zh-tw'] = dict(
    form_no='表格編號 PRI-FRM-TC',
    title_main='研究員申請表', title_sub='Research Fellow Application Form',
    cover_desc='適用於申請全景研究院研究員、副研究員、研究助理及相關學術任命。',
    issued_by_label='發布機構', issued_by_value='全景研究院 · 全景學術集團',
    appno_label='申請編號（內部填寫）',
    official_label='官方範本', footer_line='全景研究院  ·  research.panorama-sg.com  ·  官方範本',
    instructions_label='填寫說明',
    p1_label='第一部分', p1_title='一、申請說明',
    p1_body1='全景研究院是全景學術集團旗下的研究與學術發展平台，主要承擔學術研究、研究專案組織、學術出版研究、索引與評價研究、政策與社會研究、人工智慧與未來社會研究、學術會議、智庫報告及國際合作等相關工作。',
    p1_body2='本申請表用於申請研究院研究員、副研究員、研究助理等學術任命。申請人應確保所提交資訊和材料真實、準確、完整。研究院將根據申請人的教育背景、研究經歷、科研成果、研究方向匹配度、未來研究計劃及可參與研究院工作的情況進行綜合評估。',
    p1_body3='提交本申請表並不當然構成錄用、聘用、勞動關係、僱傭關係或薪酬支付承諾。具體任命結果、任期、職責、權益、合作方式及費用安排，以研究院正式書面通知、任命文件、合作協議或專案協議為準。',
    p1_note='申請編號由研究院內部受理後統一填寫，申請人無需填寫。',
    p2_label='第二部分', p2_title='二、受理資訊',
    f_appno='申請編號', f_appdate='申請日期', f_y='年', f_m='月', f_d='日',
    cat_label='申請類別',
    cat_opts=['研究員', '副研究員', '研究助理', '青年研究員', '訪問研究員', '兼職研究員', '特聘研究員', '其他'],
    term_label='擬申請任期', term_opts=['1 年', '2 年', '3 年', '其他'],
    mode_label='擬參與方式', mode_opts=['遠距參與', '專案合作', '學術任命', '線上線下結合', '其他'],
    p3_label='第三部分', p3_title='三、申請人基本資訊',
    f_fullname='中文姓名', f_nameLatin='英文姓名 / 拼音', f_nationality='國籍 / 地區', f_residence='現居國家 / 地區',
    f_tel='聯絡電話', f_email='電子郵箱', f_institution='現任單位', f_faculty='所屬院系 / 部門',
    f_position='現任職務 / 職稱', f_degree='最高學歷 / 學位', f_researchareas='主要研究方向',
    f_orcid='ORCID', f_scholar='Google Scholar / 學術主頁', f_website='個人網站 / 主頁', f_address='通訊地址',
    p4_label='第四部分', p4_title='四、教育背景', p4_note='請從最高學歷開始填寫。',
    edu_label='教育經歷', f_dates='起訖時間', f_institution2='院校名稱', f_country='國家 / 地區', f_field='專業 / 研究方向', f_degree2='學位',
    p5_label='第五部分', p5_title='五、工作經歷與學術任職',
    pos_label='經歷', f_org='單位 / 機構', f_role='職務 / 身份', f_mainresp='主要職責或學術工作：',
    p6_label='第六部分', p6_title='六、擬申請研究方向',
    center_label='擬申請加入的研究方向或研究中心',
    center_opts=['學術出版研究', '期刊索引與評價研究', '公共政策與社會研究', '人工智慧與未來社會', '教育與學習研究', '藝術、文化與具身研究', '國際關係與區域研究', '數位治理與平台治理', '青年議題與社會發展', '其他'],
    keywords_label='研究關鍵詞',
    p6_desc='請簡要說明您的主要研究方向、研究問題、研究方法，以及與全景研究院研究平台的關聯：',
    p7_label='第七部分', p7_title='七、主要科研成果',
    p7_sub1='1. 代表性論文', p7_note1='請列出近五年或最具代表性的論文，包括作者、題名、期刊名稱、年份、卷期、頁碼、DOI 或連結。',
    p7_sub2='2. 著作、教材、譯著或研究報告',
    p7_sub3='3. 其他學術成果', p7_note3='可填寫會議論文、資料庫、指數體系、課程建設、軟體工具、政策建議、媒體傳播、社會服務或知識轉化成果。',
    p8_label='第八部分', p8_title='八、科研專案與學術影響',
    proj_label='科研專案', f_projtitle='專案名稱', f_funder='資助 / 委託單位', f_yourrole='本人角色',
    status_label='專案狀態：', status_opts=['在研', '已結案', '申報中', '其他'],
    honours_label='學術榮譽、獎勵、社會影響或學術服務經歷：',
    p9_label='第九部分', p9_title='九、申請陳述',
    p9_sub1='1. 申請加入研究院的原因', p9_note1='請說明申請動機、與研究院定位的契合度，以及希望透過研究院平台開展的學術工作。',
    p9_sub2='2. 未來 1—3 年研究計劃', p9_note2='請說明未來研究主題、預期成果、論文計劃、專案計劃、報告計劃及與研究院平台的結合方式。',
    p9_sub3='3. 可為研究院提供的學術貢獻', p9_note3='可包括論文發表、課題申報、研究報告、政策簡報、學術會議、國際合作、期刊建設、資料庫建設、課程開發、學術傳播等。',
    p10_label='第十部分', p10_title='十、參與承諾與時間安排',
    time_label='預計每月可投入時間', time_opts=['1—3 小時', '4—8 小時', '9–15 小時', '15 小時以上', '其他'],
    activities_label='可參與的研究院事務',
    activities_opts=['研究專案', '研究報告 / 政策簡報', '工作論文', '論文集 / 專著', '學術會議 / 研討會', '講座 / 培訓課程', '期刊建設 / 審稿', '資料庫 / 指數建設', '國際合作', '其他'],
    commit_label='年度最低貢獻意向',
    commit_opts=['每年至少提交 1 項研究成果', '每年至少參與 1 次研究院活動', '每年至少參與 1 個研究專案', '根據研究院安排協商確定', '其他'],
    p11_label='第十一部分', p11_title='十一、語言能力與專業技能',
    lang_label='語言能力', lang_opts=['中文', '英文', '德文', '韓文', '日文', '其他'],
    skills_label='研究方法與技術能力',
    skills_opts=['定量研究', '定性研究', '混合研究', '比較研究', '案例研究', '文獻計量', '文本分析', '政策分析', 'SPSS / Stata', 'R / Python', 'LaTeX', 'NVivo', 'GIS', '資料庫建設', 'AI 工具', '網站 / 數位平台', '其他'],
    p12_label='第十二部分', p12_title='十二、學術倫理與利益衝突聲明',
    eth_q1='是否存在正在處理中的學術不端、出版倫理或研究倫理問題：',
    eth_q2='是否存在可能影響研究院任命、專案合作或學術判斷的利益衝突：',
    eth_q3='是否願意遵守研究院章程、學術規範、研究倫理、出版倫理、資料與文件規範、利益衝突政策及名稱標誌使用規範：',
    yes_no=['否', '是，請說明'], yes_no2=['是', '否'],
    p13_label='第十三部分', p13_title='十三、申請材料清單',
    req_label='必交材料', req_opts=['研究員申請表', '個人簡歷', '最高學位證明或在讀證明', '代表性論文或研究成果', '研究計劃或申請陳述', '學術主頁 / 出版物清單'],
    opt_label='選交材料', opt_opts=['推薦信', '獲獎證明', '專案證明', '任職或在職證明', '出版合約 / 發表證明', '其他補充材料'],
    p14_label='第十四部分', p14_title='十四、申請人聲明',
    decl_intro='本人鄭重聲明：',
    decl_items=[
        '本申請表及所附材料均由本人提交，內容真實、準確、完整。',
        '本人瞭解全景研究院有權依據申請材料進行資格審查、學術評估、研究方向匹配及任職審核，並可要求本人補充相關證明材料。',
        '本人瞭解全景研究院是全景學術集團旗下研究與學術發展平台，並非獨立法人實體。',
        '本人瞭解提交本表或通過審核，並不當然構成勞動關係、僱傭關係、薪酬支付承諾或其他法律意義上的聘用關係。',
        '本人承諾遵守研究院章程、學術規範、研究倫理、出版倫理、資料與文件規範、利益衝突政策及名稱標誌使用規範。',
        '本人承諾在參與研究院相關專案、出版、會議、報告、資料庫、指數建設及合作活動時，遵守學術誠信原則，不存在抄襲、偽造、篡改、重複發表、不當署名、虛假同儕審查等學術不端行為。',
        '本人同意研究院在申請審核、學術評估、任職管理、學術檔案建立、平台展示及後續聯絡過程中合理使用本人提交的資訊與材料。',
        '本人瞭解最終任命結果、任期、職責、權益及合作方式，以研究院正式書面通知、任命文件、合作協議或專案協議為準。',
    ],
    sig_label='申請人簽名', date_label='日期',
    internal_title='研究院內部審核部分', internal_subtitle='INTERNAL REVIEW ONLY · 以下內容由全景研究院內部填寫，申請人無需填寫。',
    p15_label='', p15_title='十五、初審意見',
    completeness_label='材料完整性', completeness_opts=['完整', '基本完整', '需補充', '不完整'],
    eligibility_label='申請資格', eligibility_opts=['符合', '基本符合', '需進一步核實', '不符合'],
    background_label='學術背景', background_opts=['優秀', '良好', '一般', '不符合要求'],
    alignment_label='研究方向匹配度', alignment_opts=['高度匹配', '較匹配', '一般', '不匹配'],
    prelim_comments_label='初審意見：', prelim_supplement_label='需補充材料或修改事項：',
    prelim_conclusion_label='初審結論', prelim_conclusion_opts=['通過', '有條件通過', '暫緩', '不通過'],
    reviewer_label='初審人',
    p16_label='內部審核 · 續', p16_title='十六、學術複審意見',
    committee_label='學術委員會 / 研究負責人意見：',
    recommend_label='學術任命建議',
    recommend_opts=['建議任命為研究員', '建議任命為副研究員', '建議任命為研究助理', '建議任命為青年研究員', '建議轉為訪問研究員申請', '建議暫緩', '不建議任命'],
    recterm_label='建議任期', recterm_opts=['1 年', '2 年', '3 年', '其他'],
    reccenter_label='建議歸屬研究方向 / 研究中心', reccenter_opts=['學術出版研究', '期刊索引與評價研究', '公共政策與社會研究', '人工智慧與未來社會', '藝術、文化與具身研究', '暫不指定'],
    finaldecision_label='最終審核結論', finaldecision_opts=['批准', '有條件批准', '暫緩', '不批准'],
    academiclead_label='學術負責人簽名', institutedirector_label='研究院負責人簽名',
    closing_tagline='全景學術集團旗下研究與學術發展平台',
    closing_contact='Email: research@panorama-sg.com  ·  Website: research.panorama-sg.com',
)


def build_doc(lang):
    ctx = Ctx(lang)
    t = T[lang]
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = ctx.sans
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), ctx.ea_sans)

    build_header_footer(doc, ctx, t['official_label'], t['footer_line'], t['instructions_label'])
    add_cover(doc, ctx, t['form_no'], t['title_main'], t['title_sub'], t['cover_desc'], t['issued_by_label'], t['issued_by_value'], t['appno_label'])
    doc.add_page_break()

    # Part 1
    part_heading(doc, ctx, t['p1_label'], t['p1_title'])
    body_text(doc, ctx, t['p1_body1'])
    body_text(doc, ctx, t['p1_body2'])
    body_text(doc, ctx, t['p1_body3'])
    note_box(doc, ctx, t['p1_note'])
    doc.add_page_break()

    # Part 2
    part_heading(doc, ctx, t['p2_label'], t['p2_title'])
    field_single(doc, ctx, t['f_appno'])
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t['f_appdate']); set_font(r, ctx.mono, 8.5, color=TEXT_MUTED, east_asian=ctx.ea_mono)
    dt = doc.add_paragraph(); dt.paragraph_format.space_after = Pt(10)
    r = dt.add_run(f"{t['f_y']}          {t['f_m']}          {t['f_d']}"); set_font(r, ctx.sans, 9, color=TEXT_MUTED, east_asian=ctx.ea_sans)
    add_bottom_rule(dt, color=BORDER_LIGHT, sz=6, space=1)
    sub_heading(doc, ctx, t['cat_label'])
    checkbox_grid(doc, ctx, t['cat_opts'], cols=3)
    sub_heading(doc, ctx, t['term_label'])
    checkbox_grid(doc, ctx, t['term_opts'], cols=4)
    sub_heading(doc, ctx, t['mode_label'])
    checkbox_grid(doc, ctx, t['mode_opts'], cols=2)
    doc.add_page_break()

    # Part 3
    part_heading(doc, ctx, t['p3_label'], t['p3_title'])
    field_pair(doc, ctx, t['f_fullname'], t['f_nameLatin'])
    field_pair(doc, ctx, t['f_nationality'], t['f_residence'])
    field_pair(doc, ctx, t['f_tel'], t['f_email'])
    field_pair(doc, ctx, t['f_institution'], t['f_faculty'])
    field_pair(doc, ctx, t['f_position'], t['f_degree'])
    field_single(doc, ctx, t['f_researchareas'])
    field_pair(doc, ctx, t['f_orcid'], t['f_scholar'])
    field_single(doc, ctx, t['f_website'])
    field_single(doc, ctx, t['f_address'])
    doc.add_page_break()

    # Part 4
    part_heading(doc, ctx, t['p4_label'], t['p4_title'])
    body_text(doc, ctx, t['p4_note'])
    for i in (1, 2, 3):
        def body(cell, i=i):
            fp = cell.add_table(rows=1, cols=2)
            no_borders(fp.rows[0].cells[0]); no_borders(fp.rows[0].cells[1])
            _field_cell(fp.rows[0].cells[0], ctx, t['f_dates'])
            _field_cell(fp.rows[0].cells[1], ctx, t['f_institution2'])
            fp2 = cell.add_table(rows=1, cols=2)
            no_borders(fp2.rows[0].cells[0]); no_borders(fp2.rows[0].cells[1])
            _field_cell(fp2.rows[0].cells[0], ctx, t['f_country'])
            _field_cell(fp2.rows[0].cells[1], ctx, t['f_field'])
            fs = cell.add_paragraph(); fs.paragraph_format.space_after = Pt(2)
            r = fs.add_run(t['f_degree2']); set_font(r, ctx.mono, 8.5, color=TEXT_MUTED, east_asian=ctx.ea_mono)
            fb = cell.add_paragraph(); fb.paragraph_format.space_after = Pt(8)
            add_bottom_rule(fb, color=BORDER_LIGHT, sz=6, space=1)
            rr = fb.add_run(' '); set_font(rr, ctx.sans, 8)
        boxed_repeat(doc, ctx, f"{t['edu_label']} {i}", body)
    doc.add_page_break()

    # Part 5
    part_heading(doc, ctx, t['p5_label'], t['p5_title'])
    for i in (1, 2, 3):
        def body(cell, i=i):
            fp = cell.add_table(rows=1, cols=2)
            no_borders(fp.rows[0].cells[0]); no_borders(fp.rows[0].cells[1])
            _field_cell(fp.rows[0].cells[0], ctx, t['f_dates'])
            _field_cell(fp.rows[0].cells[1], ctx, t['f_org'])
            fs = cell.add_paragraph(); fs.paragraph_format.space_after = Pt(2)
            r = fs.add_run(t['f_role']); set_font(r, ctx.mono, 8.5, color=TEXT_MUTED, east_asian=ctx.ea_mono)
            fb = cell.add_paragraph(); fb.paragraph_format.space_after = Pt(8)
            add_bottom_rule(fb, color=BORDER_LIGHT, sz=6, space=1)
            rr = fb.add_run(' '); set_font(rr, ctx.sans, 8)
            mr = cell.add_paragraph(); mr.paragraph_format.space_after = Pt(4)
            r = mr.add_run(t['f_mainresp']); set_font(r, ctx.sans, 9.5, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)
            for _ in range(2):
                lp = cell.add_paragraph(); lp.paragraph_format.space_after = Pt(8)
                add_bottom_rule(lp, color=BORDER_LIGHT, sz=6, space=1)
                rr = lp.add_run(' '); set_font(rr, ctx.sans, 8)
        boxed_repeat(doc, ctx, f"{t['pos_label']} {i}", body)
    doc.add_page_break()

    # Part 6
    part_heading(doc, ctx, t['p6_label'], t['p6_title'])
    sub_heading(doc, ctx, t['center_label'])
    checkbox_grid(doc, ctx, t['center_opts'], cols=2)
    sub_heading(doc, ctx, t['keywords_label'])
    numbered_lines(doc, ctx, 1, 5)
    body_text(doc, ctx, t['p6_desc'])
    lines(doc, ctx, 3)
    doc.add_page_break()

    # Part 7
    part_heading(doc, ctx, t['p7_label'], t['p7_title'])
    sub_heading(doc, ctx, t['p7_sub1'], t['p7_note1'])
    numbered_lines(doc, ctx, 1, 5)
    sub_heading(doc, ctx, t['p7_sub2'])
    numbered_lines(doc, ctx, 1, 3)
    sub_heading(doc, ctx, t['p7_sub3'], t['p7_note3'])
    lines(doc, ctx, 2)
    doc.add_page_break()

    # Part 8
    part_heading(doc, ctx, t['p8_label'], t['p8_title'])
    for i in (1, 2, 3):
        def body(cell, i=i):
            fp = cell.add_table(rows=1, cols=2)
            no_borders(fp.rows[0].cells[0]); no_borders(fp.rows[0].cells[1])
            _field_cell(fp.rows[0].cells[0], ctx, t['f_projtitle'])
            _field_cell(fp.rows[0].cells[1], ctx, t['f_funder'])
            fp2 = cell.add_table(rows=1, cols=2)
            no_borders(fp2.rows[0].cells[0]); no_borders(fp2.rows[0].cells[1])
            _field_cell(fp2.rows[0].cells[0], ctx, t['f_yourrole'])
            _field_cell(fp2.rows[0].cells[1], ctx, t['f_dates'])
            sp = cell.add_paragraph(); sp.paragraph_format.space_after = Pt(6)
            r = sp.add_run(t['status_label'] + '  ')
            set_font(r, ctx.sans, 9.5, bold=True, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)
            for opt in t['status_opts']:
                r = sp.add_run(CHECK + '  ' + opt + '   ')
                set_font(r, ctx.sans, 9.5, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)
        boxed_repeat(doc, ctx, f"{t['proj_label']} {i}", body)
    body_text(doc, ctx, t['honours_label'])
    lines(doc, ctx, 2)
    doc.add_page_break()

    # Part 9
    part_heading(doc, ctx, t['p9_label'], t['p9_title'])
    sub_heading(doc, ctx, t['p9_sub1'], t['p9_note1'])
    lines(doc, ctx, 2)
    sub_heading(doc, ctx, t['p9_sub2'], t['p9_note2'])
    lines(doc, ctx, 2)
    sub_heading(doc, ctx, t['p9_sub3'], t['p9_note3'])
    lines(doc, ctx, 2)
    doc.add_page_break()

    # Part 10
    part_heading(doc, ctx, t['p10_label'], t['p10_title'])
    sub_heading(doc, ctx, t['time_label'])
    checkbox_grid(doc, ctx, t['time_opts'], cols=3)
    sub_heading(doc, ctx, t['activities_label'])
    checkbox_grid(doc, ctx, t['activities_opts'], cols=2)
    sub_heading(doc, ctx, t['commit_label'])
    checkbox_grid(doc, ctx, t['commit_opts'], cols=1)
    doc.add_page_break()

    # Part 11
    part_heading(doc, ctx, t['p11_label'], t['p11_title'])
    sub_heading(doc, ctx, t['lang_label'])
    checkbox_grid(doc, ctx, t['lang_opts'], cols=3)
    sub_heading(doc, ctx, t['skills_label'])
    checkbox_grid(doc, ctx, t['skills_opts'], cols=3)
    doc.add_page_break()

    # Part 12
    part_heading(doc, ctx, t['p12_label'], t['p12_title'])
    checkbox_inline(doc, ctx, t['eth_q1'], t['yes_no'])
    checkbox_inline(doc, ctx, t['eth_q2'], t['yes_no'])
    checkbox_inline(doc, ctx, t['eth_q3'], t['yes_no2'])
    doc.add_page_break()

    # Part 13
    part_heading(doc, ctx, t['p13_label'], t['p13_title'])
    sub_heading(doc, ctx, t['req_label'])
    checkbox_grid(doc, ctx, t['req_opts'], cols=2)
    sub_heading(doc, ctx, t['opt_label'])
    checkbox_grid(doc, ctx, t['opt_opts'], cols=2)
    doc.add_page_break()

    # Part 14
    part_heading(doc, ctx, t['p14_label'], t['p14_title'])
    body_text(doc, ctx, t['decl_intro'])
    for i, item in enumerate(t['decl_items'], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(f'{i}.  ')
        set_font(r, ctx.sans, 10, bold=True, color=TEXT_MUTED, east_asian=ctx.ea_sans)
        r = p.add_run(item)
        set_font(r, ctx.sans, 10, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    signature_row(doc, ctx, t['sig_label'], t['date_label'], t['f_y'], t['f_m'], t['f_d'])
    doc.add_page_break()

    # Internal review
    dark_banner(doc, ctx, t['internal_title'], t['internal_subtitle'])
    part_heading(doc, ctx, t['p15_label'], t['p15_title'])
    table = doc.add_table(rows=1, cols=2)
    ca, cb = table.rows[0].cells
    no_borders(ca); no_borders(cb)
    p = ca.paragraphs[0]; r = p.add_run(t['completeness_label']); set_font(r, ctx.sans, 10, bold=True, color=TEXT_MAIN, east_asian=ctx.ea_sans)
    for opt in t['completeness_opts']:
        pp = ca.add_paragraph(); pp.paragraph_format.space_after = Pt(3)
        r = pp.add_run(CHECK + '  ' + opt); set_font(r, ctx.sans, 9.5, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)
    p = cb.paragraphs[0]; r = p.add_run(t['eligibility_label']); set_font(r, ctx.sans, 10, bold=True, color=TEXT_MAIN, east_asian=ctx.ea_sans)
    for opt in t['eligibility_opts']:
        pp = cb.add_paragraph(); pp.paragraph_format.space_after = Pt(3)
        r = pp.add_run(CHECK + '  ' + opt); set_font(r, ctx.sans, 9.5, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    table2 = doc.add_table(rows=1, cols=2)
    ca, cb = table2.rows[0].cells
    no_borders(ca); no_borders(cb)
    p = ca.paragraphs[0]; r = p.add_run(t['background_label']); set_font(r, ctx.sans, 10, bold=True, color=TEXT_MAIN, east_asian=ctx.ea_sans)
    for opt in t['background_opts']:
        pp = ca.add_paragraph(); pp.paragraph_format.space_after = Pt(3)
        r = pp.add_run(CHECK + '  ' + opt); set_font(r, ctx.sans, 9.5, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)
    p = cb.paragraphs[0]; r = p.add_run(t['alignment_label']); set_font(r, ctx.sans, 10, bold=True, color=TEXT_MAIN, east_asian=ctx.ea_sans)
    for opt in t['alignment_opts']:
        pp = cb.add_paragraph(); pp.paragraph_format.space_after = Pt(3)
        r = pp.add_run(CHECK + '  ' + opt); set_font(r, ctx.sans, 9.5, color=TEXT_SECONDARY, east_asian=ctx.ea_sans)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    body_text(doc, ctx, t['prelim_comments_label'], space_after=4)
    lines(doc, ctx, 2)
    body_text(doc, ctx, t['prelim_supplement_label'], space_after=4)
    lines(doc, ctx, 1)
    sub_heading(doc, ctx, t['prelim_conclusion_label'])
    checkbox_grid(doc, ctx, t['prelim_conclusion_opts'], cols=4)
    field_single(doc, ctx, t['reviewer_label'])
    doc.add_page_break()

    # Part 16
    part_heading(doc, ctx, t['p16_label'], t['p16_title'])
    body_text(doc, ctx, t['committee_label'], space_after=4)
    lines(doc, ctx, 2)
    sub_heading(doc, ctx, t['recommend_label'])
    checkbox_grid(doc, ctx, t['recommend_opts'], cols=2)
    sub_heading(doc, ctx, t['recterm_label'])
    checkbox_grid(doc, ctx, t['recterm_opts'], cols=4)
    sub_heading(doc, ctx, t['reccenter_label'])
    checkbox_grid(doc, ctx, t['reccenter_opts'], cols=2)
    sub_heading(doc, ctx, t['finaldecision_label'])
    checkbox_grid(doc, ctx, t['finaldecision_opts'], cols=4)
    signature_row(doc, ctx, t['academiclead_label'], t['date_label'], t['f_y'], t['f_m'], t['f_d'])
    signature_row(doc, ctx, t['institutedirector_label'], t['date_label'], t['f_y'], t['f_m'], t['f_d'])

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    closing = doc.add_table(rows=1, cols=2)
    ca, cb = closing.rows[0].cells
    no_borders(ca); no_borders(cb)
    p = ca.paragraphs[0]
    run = p.add_run()
    run.add_picture(LOGO_HEADER, width=Cm(2.6))
    p2 = cb.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p2.add_run(t['closing_tagline'])
    set_font(r, ctx.sans, 9, color=TEXT_MUTED, east_asian=ctx.ea_sans)
    p3 = cb.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p3.add_run(t['closing_contact'])
    set_font(r, ctx.mono, 9, color=TEXT_MUTED, east_asian=ctx.ea_mono)

    return doc


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    if not os.path.exists(LOGO_HEADER) or not os.path.exists(LOGO_COVER):
        raise SystemExit(
            'Logo assets not found. Expected:\n  ' + LOGO_HEADER + '\n  ' + LOGO_COVER +
            '\n\nSet PRI_LOGO_ASSETS_DIR to the directory containing logo-header.png and logo-cover.png.'
        )
    for lang in ('en', 'zh-cn', 'zh-tw'):
        doc = build_doc(lang)
        filename = f"PRI_Research Fellow Application Form_{LANG_SUFFIX[lang]}.docx"
        out_path = os.path.join(OUT_DIR, filename)
        doc.save(out_path)
        print(f'Wrote {filename}')
    print(f'\nGenerated 3 application form documents in {OUT_DIR}')


if __name__ == '__main__':
    main()
