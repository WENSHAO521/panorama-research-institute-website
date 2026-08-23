# One-time content-generation script: builds the six official .docx
# document templates listed on /resources/templates (Research Proposal,
# Research Report, Policy Brief, Working Paper, Event Minutes Summary,
# Edited Volume Chapter), in English, Simplified Chinese, and Traditional
# Chinese, and writes them into public/templates/.
#
# Usage: python scripts/generate-document-templates.py
#
# Requires: python-docx  (pip install python-docx)
# Uses the official brand mark rasterized from public/brand/logo-mono-black.svg
# and public/brand/logo-stacked.svg (via `sharp`, at build time — see the
# "Regenerating the logo PNGs" note at the bottom of this file). The
# rasterized PNGs are build-time intermediates, not checked into the repo;
# point LOGO_HEADER / LOGO_COVER at wherever you generated them.

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(ROOT, 'public', 'templates')
LOGO_ASSETS_DIR = os.environ.get('PRI_LOGO_ASSETS_DIR', os.path.join(ROOT, '.logo-assets'))
LOGO_HEADER = os.path.join(LOGO_ASSETS_DIR, 'logo-header.png')
LOGO_COVER = os.path.join(LOGO_ASSETS_DIR, 'logo-cover.png')

TEXT_MAIN = RGBColor(0x1F, 0x23, 0x26)
TEXT_SECONDARY = RGBColor(0x5C, 0x63, 0x68)
TEXT_MUTED = RGBColor(0x8A, 0x91, 0x96)
ACCENT_STEEL = RGBColor(0x37, 0x5A, 0x6B)
BORDER_LIGHT = 'DADDDD'
BG_SOFT = 'F1F2F2'

LANGS = ['en', 'zh-cn', 'zh-tw']
LANG_SUFFIX = {'en': 'EN', 'zh-cn': 'CN', 'zh-tw': 'CN-TW'}

# East-Asian companion fonts. Word's automatic CJK fallback (when only a
# Latin font like Calibri is specified) is inconsistent about which glyphs
# get the requested italic slant, so zh-cn/zh-tw documents specify real CJK
# fonts explicitly instead of relying on fallback.
CJK_SANS = {'en': 'Calibri', 'zh-cn': 'Microsoft YaHei', 'zh-tw': 'Microsoft JhengHei'}
CJK_SERIF = {'en': 'Georgia', 'zh-cn': 'SimSun', 'zh-tw': 'PMingLiU'}
CJK_MONO = {'en': 'Consolas', 'zh-cn': 'Microsoft YaHei', 'zh-tw': 'Microsoft JhengHei'}

L = {
    'official_template': {'en': 'Official Template', 'zh-cn': '官方模板', 'zh-tw': '官方範本'},
    'template_details': {'en': 'Template Details', 'zh-cn': '模板信息', 'zh-tw': '範本資訊'},
    'template_structure': {'en': 'Template Structure', 'zh-cn': '模板结构', 'zh-tw': '範本結構'},
    'format': {'en': 'Format', 'zh-cn': '格式', 'zh-tw': '格式'},
    'format_v': {'en': 'Microsoft Word (.docx)', 'zh-cn': 'Microsoft Word（.docx）', 'zh-tw': 'Microsoft Word（.docx）'},
    'version': {'en': 'Version', 'zh-cn': '版本', 'zh-tw': '版本'},
    'version_v': {'en': 'v1.0 (2026)', 'zh-cn': 'v1.0（2026）', 'zh-tw': 'v1.0（2026）'},
    'submission_lang': {'en': 'Language of Submission', 'zh-cn': '提交语言', 'zh-tw': '提交語言'},
    'submission_lang_v': {'en': 'English', 'zh-cn': '英文', 'zh-tw': '英文'},
    'instructions': {'en': 'Instructions', 'zh-cn': '填写说明', 'zh-tw': '填寫說明'},
    'footer_line': {
        'en': 'Panorama Research Institute  ·  research.panorama-sg.com  ·  Official Template',
        'zh-cn': '全景研究院  ·  research.panorama-sg.com  ·  官方模板',
        'zh-tw': '全景研究院  ·  research.panorama-sg.com  ·  官方範本',
    },
    'do_not_alter': {
        'en': 'Standard Institute text — do not alter.',
        'zh-cn': '研究院标准文本 — 请勿修改。',
        'zh-tw': '研究院標準文本 — 請勿修改。',
    },
    'legal_note': {
        'en': 'Panorama Research Institute is an internal research and academic development division of Panorama Scholarly Group and is not a separate legal entity.',
        'zh-cn': '全景研究院是全景学术集团的内部研究与学术发展部门，不具备独立法人资格。',
        'zh-tw': '全景研究院是全景學術集團的內部研究與學術發展部門，不具備獨立法人資格。',
    },
    'about_institute_body': {
        'en': 'Panorama Research Institute is an internal research and academic development division of Panorama Scholarly Group. It maintains academic independence in its research activities and is not a separate legal entity. For more information, visit research.panorama-sg.com.',
        'zh-cn': '全景研究院是全景学术集团的内部研究与学术发展部门。研究院在其研究活动中保持学术独立，不具备独立法人资格。欲了解更多信息，请访问 research.panorama-sg.com。',
        'zh-tw': '全景研究院是全景學術集團的內部研究與學術發展部門。研究院在其研究活動中保持學術獨立，不具備獨立法人資格。欲了解更多資訊，請造訪 research.panorama-sg.com。',
    },
}


def set_font(run, font_name, size=None, bold=None, italic=None, color=None, east_asian=None):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)
    rfonts.set(qn('w:eastAsia'), east_asian or font_name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, color=BORDER_LIGHT, sz=4, style='dashed'):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), style)
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        borders.append(el)
    tc_pr.append(borders)


def add_bottom_rule(paragraph, color=BORDER_LIGHT, sz=6):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    set_font(run, 'Consolas', 9, color=TEXT_MUTED)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def build_header_footer(doc, lang, sans, mono, ea_sans, ea_mono):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    header = section.header
    h_table = header.add_table(rows=1, cols=2, width=Cm(16.0))
    h_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    h_table.autofit = True
    left_cell, right_cell = h_table.rows[0].cells
    left_p = left_cell.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = left_p.add_run()
    run.add_picture(LOGO_HEADER, width=Cm(3.6))
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = right_p.add_run(L['official_template'][lang])
    set_font(r, mono, 9, color=TEXT_MUTED, east_asian=ea_mono)
    add_bottom_rule(header.paragraphs[0] if header.paragraphs[0].runs == [] else right_p)
    # thin rule under the whole header block
    rule_p = header.add_paragraph()
    add_bottom_rule(rule_p, color=BORDER_LIGHT, sz=6)

    footer = section.footer
    f_table = footer.add_table(rows=1, cols=2, width=Cm(16.0))
    f_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fl_cell, fr_cell = f_table.rows[0].cells
    fl_p = fl_cell.paragraphs[0]
    r = fl_p.add_run(L['footer_line'][lang])
    set_font(r, mono, 8.5, color=TEXT_MUTED, east_asian=ea_mono)
    fr_p = fr_cell.paragraphs[0]
    fr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fr_p.add_run(L['instructions'][lang] + ' · ')
    set_font(r, mono, 8.5, color=TEXT_MUTED, east_asian=ea_mono)
    add_page_number_field(fr_p)


def add_cover(doc, lang, title, subtitle, sans, serif, mono, ea_sans, ea_serif, ea_mono):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(LOGO_COVER, width=Cm(4.6))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_font(r, serif, 26, bold=False, color=TEXT_MAIN, east_asian=ea_serif)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run(subtitle)
    set_font(r, mono, 10.5, color=ACCENT_STEEL, east_asian=ea_mono)

    rule = doc.add_paragraph()
    add_bottom_rule(rule, color=BORDER_LIGHT, sz=8)
    rule.paragraph_format.space_after = Pt(22)


def add_meta_table(doc, lang, rows, sans, mono, ea_sans, ea_mono):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(5.2)
    table.columns[1].width = Cm(10.8)
    for label, value in rows:
        row = table.add_row()
        lcell, vcell = row.cells
        shade_cell(lcell, BG_SOFT)
        shade_cell(vcell, 'FFFFFF')
        set_cell_border(lcell, style='single', color=BORDER_LIGHT, sz=4)
        set_cell_border(vcell, style='single', color=BORDER_LIGHT, sz=4)
        lp = lcell.paragraphs[0]
        lp.paragraph_format.space_before = Pt(4)
        lp.paragraph_format.space_after = Pt(4)
        r = lp.add_run(label)
        set_font(r, mono, 9.5, bold=True, color=TEXT_SECONDARY, east_asian=ea_mono)
        vp = vcell.paragraphs[0]
        vp.paragraph_format.space_before = Pt(4)
        vp.paragraph_format.space_after = Pt(4)
        r = vp.add_run(value)
        set_font(r, sans, 10.5, color=TEXT_MAIN, east_asian=ea_sans)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_intro(doc, text, sans, ea_sans):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run(text)
    set_font(r, sans, 11, color=TEXT_SECONDARY, east_asian=ea_sans)


def add_section(doc, lang, heading, note, placeholder, sans, serif, mono, ea_sans, ea_serif):
    hp = doc.add_paragraph()
    hp.paragraph_format.space_before = Pt(16)
    hp.paragraph_format.space_after = Pt(4)
    r = hp.add_run(heading)
    set_font(r, serif, 14, bold=True, color=TEXT_MAIN, east_asian=ea_serif)

    if note:
        np = doc.add_paragraph()
        np.paragraph_format.space_after = Pt(8)
        r = np.add_run(note)
        set_font(r, sans, 9.5, italic=True, color=TEXT_MUTED, east_asian=ea_sans)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    shade_cell(cell, BG_SOFT)
    set_cell_border(cell, color=BORDER_LIGHT, sz=4, style='dashed')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(placeholder)
    set_font(r, sans, 10.5, italic=True, color=TEXT_MUTED, east_asian=ea_sans)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_doc(lang, title, subtitle, meta_rows, intro, sections):
    sans = 'Calibri'
    serif = 'Georgia'
    mono = 'Consolas'
    ea_sans = CJK_SANS[lang]
    ea_serif = CJK_SERIF[lang]
    ea_mono = CJK_MONO[lang]

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = sans
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), ea_sans)

    build_header_footer(doc, lang, sans, mono, ea_sans, ea_mono)
    add_cover(doc, lang, title, subtitle, sans, serif, mono, ea_sans, ea_serif, ea_mono)

    dh = doc.add_paragraph()
    dh.paragraph_format.space_after = Pt(8)
    r = dh.add_run(L['template_details'][lang])
    set_font(r, mono, 10, bold=True, color=TEXT_SECONDARY, east_asian=ea_mono)
    add_meta_table(doc, lang, meta_rows, sans, mono, ea_sans, ea_mono)

    add_intro(doc, intro, sans, ea_sans)

    sh = doc.add_paragraph()
    sh.paragraph_format.space_before = Pt(6)
    sh.paragraph_format.space_after = Pt(4)
    r = sh.add_run(L['template_structure'][lang])
    set_font(r, serif, 17, bold=False, color=TEXT_MAIN, east_asian=ea_serif)
    rule = doc.add_paragraph()
    add_bottom_rule(rule, color=BORDER_LIGHT, sz=6)
    rule.paragraph_format.space_after = Pt(10)

    for heading, note, placeholder in sections:
        add_section(doc, lang, heading, note, placeholder, sans, serif, mono, ea_sans, ea_serif)

    return doc


# ---------------------------------------------------------------------------
# Template content
# ---------------------------------------------------------------------------

TEMPLATES = []

# 1. Research Proposal Template ---------------------------------------------
TEMPLATES.append({
    'slug': 'Research Proposal Template',
    'title': {'en': 'Research Proposal Template', 'zh-cn': '研究提案模板', 'zh-tw': '研究提案範本'},
    'subtitle': {
        'en': 'Official Template · v1.0 (2026)',
        'zh-cn': '官方模板 · v1.0（2026）',
        'zh-tw': '官方範本 · v1.0（2026）',
    },
    'meta': lambda lang: [
        (L['format'][lang], L['format_v'][lang]),
        (L['version'][lang], L['version_v'][lang]),
        ({'en': 'Access', 'zh-cn': '获取方式', 'zh-tw': '取得方式'}[lang],
         {'en': 'Open — no registration required', 'zh-cn': '开放获取 — 无需注册', 'zh-tw': '開放取得 — 無需註冊'}[lang]),
        ({'en': 'For Use With', 'zh-cn': '适用场景', 'zh-tw': '適用場景'}[lang],
         {'en': 'Research Project Proposal submissions', 'zh-cn': '提交研究项目提案', 'zh-tw': '提交研究項目提案'}[lang]),
    ],
    'intro': {
        'en': 'The Research Proposal Template provides a structured format for preparing submissions to the Institute. It includes placeholder text, formatting guidance, and section instructions. Complete all sections before submitting; incomplete proposals will be returned. Proposals may be submitted by email attachment to research@panorama-sg.com or via the Research Project Proposal web form.',
        'zh-cn': '本研究提案模板为向研究院提交提案提供了结构化格式，包含占位文字、格式说明与各章节填写指引。请在提交前完成全部章节；不完整的提案将被退回。可通过电子邮件将本文件作为附件发送至 research@panorama-sg.com，或通过网站上的研究项目提案表单提交。',
        'zh-tw': '本研究提案範本為向研究院提交提案提供了結構化格式，包含佔位文字、格式說明與各章節填寫指引。請於提交前完成全部章節；不完整的提案將被退回。可透過電子郵件將本文件作為附件傳送至 research@panorama-sg.com，或透過網站上的研究項目提案表單提交。',
    },
    'sections': lambda lang: [
        (
            {'en': 'Cover Page', 'zh-cn': '封面', 'zh-tw': '封面'}[lang],
            {'en': 'Project title, lead researcher details, proposed research center, submission date.',
             'zh-cn': '项目标题、首席研究员信息、拟属研究中心、提交日期。',
             'zh-tw': '項目標題、首席研究員資訊、擬屬研究中心、提交日期。'}[lang],
            {'en': '[Project Title]\n[Lead Researcher Name, Affiliation]\n[Proposed Research Center]\n[Submission Date]',
             'zh-cn': '［项目标题］\n［首席研究员姓名、所属机构］\n［拟属研究中心］\n［提交日期］',
             'zh-tw': '［項目標題］\n［首席研究員姓名、所屬機構］\n［擬屬研究中心］\n［提交日期］'}[lang],
        ),
        (
            {'en': '1. Project Title and Overview', 'zh-cn': '一、项目标题与概述', 'zh-tw': '一、項目標題與概述'}[lang],
            {'en': 'Concise title and a 200-word summary of the proposed project.',
             'zh-cn': '简明标题及不超过 200 字的项目摘要。',
             'zh-tw': '簡明標題及不超過 200 字的項目摘要。'}[lang],
            {'en': '[Insert a concise project title and a 200-word overview describing the aims and significance of the proposed research.]',
             'zh-cn': '［请填写简明的项目标题，并以约 200 字概述拟议研究的目标与意义。］',
             'zh-tw': '［請填寫簡明的項目標題，並以約 200 字概述擬議研究的目標與意義。］'}[lang],
        ),
        (
            {'en': '2. Research Questions', 'zh-cn': '二、研究问题', 'zh-tw': '二、研究問題'}[lang],
            {'en': 'Specific, clearly stated research questions.',
             'zh-cn': '明确、具体的研究问题。',
             'zh-tw': '明確、具體的研究問題。'}[lang],
            {'en': '[List the specific research questions this project will address, numbered.]',
             'zh-cn': '［请列出本项目拟解决的具体研究问题，并编号。］',
             'zh-tw': '［請列出本項目擬解決的具體研究問題，並編號。］'}[lang],
        ),
        (
            {'en': '3. Background and Rationale', 'zh-cn': '三、背景与研究依据', 'zh-tw': '三、背景與研究依據'}[lang],
            {'en': 'Context, literature review summary, and justification for the project.',
             'zh-cn': '研究背景、文献综述概要及项目立项依据。',
             'zh-tw': '研究背景、文獻綜述概要及項目立項依據。'}[lang],
            {'en': '[Insert the background, a brief review of relevant literature, and the rationale for undertaking this project.]',
             'zh-cn': '［请填写研究背景、相关文献综述概要，以及开展本项目的依据。］',
             'zh-tw': '［請填寫研究背景、相關文獻綜述概要，以及開展本項目的依據。］'}[lang],
        ),
        (
            {'en': '4. Methodology', 'zh-cn': '四、研究方法', 'zh-tw': '四、研究方法'}[lang],
            {'en': 'Research design, methods, and justification.',
             'zh-cn': '研究设计、研究方法及其合理性说明。',
             'zh-tw': '研究設計、研究方法及其合理性說明。'}[lang],
            {'en': '[Describe the research design and methods to be used, and justify their appropriateness for the research questions.]',
             'zh-cn': '［请描述拟采用的研究设计与方法，并说明其与研究问题的适配性。］',
             'zh-tw': '［請描述擬採用的研究設計與方法，並說明其與研究問題的適配性。］'}[lang],
        ),
        (
            {'en': '5. Expected Outputs', 'zh-cn': '五、预期成果', 'zh-tw': '五、預期成果'}[lang],
            {'en': 'Publications, datasets, or other outputs, with a timeline.',
             'zh-cn': '出版物、数据集或其他成果，并附时间安排。',
             'zh-tw': '出版物、資料集或其他成果，並附時間安排。'}[lang],
            {'en': '[List expected outputs (e.g., working paper, dataset, policy brief) and the anticipated timeline for each.]',
             'zh-cn': '［请列出预期成果（如工作论文、数据集、政策简报等）及各项成果的预计完成时间。］',
             'zh-tw': '［請列出預期成果（如工作論文、資料集、政策簡報等）及各項成果的預計完成時間。］'}[lang],
        ),
        (
            {'en': '6. Proposed Research Center', 'zh-cn': '六、拟属研究中心', 'zh-tw': '六、擬屬研究中心'}[lang],
            {'en': 'Preferred Institute research center affiliation.',
             'zh-cn': '拟隶属的研究院研究中心。',
             'zh-tw': '擬隸屬的研究院研究中心。'}[lang],
            {'en': '[State the research center this project is proposed under, or indicate "Open to recommendation."]',
             'zh-cn': '［请说明本项目拟隶属的研究中心，或填写"接受建议"。］',
             'zh-tw': '［請說明本項目擬隸屬的研究中心，或填寫「接受建議」。］'}[lang],
        ),
        (
            {'en': '7. Lead Researcher Profile', 'zh-cn': '七、首席研究员简介', 'zh-tw': '七、首席研究員簡介'}[lang],
            {'en': 'Academic bio and relevant publications.',
             'zh-cn': '学术简介及相关出版物。',
             'zh-tw': '學術簡介及相關出版物。'}[lang],
            {'en': '[Insert a brief academic biography of the lead researcher and a list of relevant publications.]',
             'zh-cn': '［请填写首席研究员的学术简介及相关出版物清单。］',
             'zh-tw': '［請填寫首席研究員的學術簡介及相關出版物清單。］'}[lang],
        ),
        (
            {'en': '8. Co-investigators (if any)', 'zh-cn': '八、共同研究者（如有）', 'zh-tw': '八、共同研究者（如有）'}[lang],
            {'en': 'Names, institutions, and roles.',
             'zh-cn': '姓名、所属机构及分工。',
             'zh-tw': '姓名、所屬機構及分工。'}[lang],
            {'en': '[List any co-investigators, their institutional affiliation, and their role in the project.]',
             'zh-cn': '［请列出共同研究者姓名、所属机构及其在项目中的分工。］',
             'zh-tw': '［請列出共同研究者姓名、所屬機構及其在項目中的分工。］'}[lang],
        ),
        (
            {'en': 'Ethics Declaration', 'zh-cn': '伦理声明', 'zh-tw': '倫理聲明'}[lang],
            {'en': 'Identification of any ethical considerations.',
             'zh-cn': '说明本项目涉及的任何伦理考量。',
             'zh-tw': '說明本項目涉及的任何倫理考量。'}[lang],
            {'en': '[Identify any ethical considerations (e.g., human subjects, sensitive data) and how they will be addressed, or state "No ethical considerations identified."]',
             'zh-cn': '［请说明本项目涉及的任何伦理考量（如涉及人类受试者、敏感数据等）及应对方式，若无则填写"未涉及伦理考量"。］',
             'zh-tw': '［請說明本項目涉及的任何倫理考量（如涉及人類受試者、敏感資料等）及因應方式，若無則填寫「未涉及倫理考量」。］'}[lang],
        ),
    ],
})

# 2. Research Report Template -------------------------------------------------
TEMPLATES.append({
    'slug': 'Research Report Template',
    'title': {'en': 'Research Report Template', 'zh-cn': '研究报告模板', 'zh-tw': '研究報告範本'},
    'subtitle': {
        'en': 'Official Template · v1.0 (2026)',
        'zh-cn': '官方模板 · v1.0（2026）',
        'zh-tw': '官方範本 · v1.0（2026）',
    },
    'meta': lambda lang: [
        (L['format'][lang], L['format_v'][lang]),
        (L['version'][lang], L['version_v'][lang]),
        ({'en': 'Typical Length', 'zh-cn': '一般篇幅', 'zh-tw': '一般篇幅'}[lang],
         {'en': '20,000–60,000 words', 'zh-cn': '20,000–60,000 字', 'zh-tw': '20,000–60,000 字'}[lang]),
        ({'en': 'Citation Style', 'zh-cn': '引用格式', 'zh-tw': '引用格式'}[lang],
         {'en': 'APA 7th (preferred) or Chicago 18th', 'zh-cn': 'APA 第 7 版（优先）或芝加哥格式第 18 版', 'zh-tw': 'APA 第 7 版（優先）或芝加哥格式第 18 版'}[lang]),
        (L['submission_lang'][lang], L['submission_lang_v'][lang]),
    ],
    'intro': {
        'en': 'Research reports must follow the Institute’s standard structure. The template provides pre-formatted sections, style definitions, and placeholder content for all required elements. Authors should not alter the section order or remove required elements. See the Citation Guide for the PSG Author-Date format used in Institute references (Research Report No. RR-YYYY-00X).',
        'zh-cn': '研究报告必须遵循研究院的标准结构。本模板提供了预先格式化的章节、样式设定以及各必需要素的占位内容。作者不得更改章节顺序或删除必需要素。研究院参考文献采用 PSG 作者-出版年格式，编号格式为"研究报告 No. RR-YYYY-00X"，具体规则请参阅引用指南。',
        'zh-tw': '研究報告必須遵循研究院的標準結構。本範本提供了預先格式化的章節、樣式設定以及各必需要素的佔位內容。作者不得更改章節順序或刪除必需要素。研究院參考文獻採用 PSG 作者-出版年格式，編號格式為「研究報告 No. RR-YYYY-00X」，具體規則請參閱引用指南。',
    },
    'sections': lambda lang: [
        (
            {'en': 'Cover Page', 'zh-cn': '封面', 'zh-tw': '封面'}[lang],
            {'en': 'Title, report number, author(s), date, research center, DOI.',
             'zh-cn': '标题、报告编号、作者、日期、研究中心、DOI。',
             'zh-tw': '標題、報告編號、作者、日期、研究中心、DOI。'}[lang],
            {'en': '[Report Title]\nResearch Report No. RR-YYYY-00X\n[Author(s), Affiliation]\n[Research Center]\n[Date]\n[DOI: 10.XXXXX/YYYYY]',
             'zh-cn': '［报告标题］\n研究报告编号 RR-YYYY-00X\n［作者、所属机构］\n［研究中心］\n［日期］\n［DOI：10.XXXXX/YYYYY］',
             'zh-tw': '［報告標題］\n研究報告編號 RR-YYYY-00X\n［作者、所屬機構］\n［研究中心］\n［日期］\n［DOI：10.XXXXX/YYYYY］'}[lang],
        ),
        (
            {'en': 'Disclaimer', 'zh-cn': '免责声明', 'zh-tw': '免責聲明'}[lang],
            L['do_not_alter'][lang],
            {'en': 'The findings, interpretations, and conclusions expressed in this report are those of the author(s) and do not necessarily reflect an official position of Panorama Research Institute. ' + L['legal_note']['en'],
             'zh-cn': '本报告所表达的研究发现、诠释与结论仅代表作者本人观点，不必然反映全景研究院的官方立场。' + L['legal_note']['zh-cn'],
             'zh-tw': '本報告所表達的研究發現、詮釋與結論僅代表作者本人觀點，不必然反映全景研究院的官方立場。' + L['legal_note']['zh-tw']}[lang],
        ),
        (
            {'en': 'Abstract', 'zh-cn': '摘要', 'zh-tw': '摘要'}[lang],
            {'en': '200–300 words; no citations in the abstract.',
             'zh-cn': '200–300 字；摘要中不得包含引用。',
             'zh-tw': '200–300 字；摘要中不得包含引用。'}[lang],
            {'en': '[Insert a 200–300 word abstract summarizing the report’s purpose, methods, and key findings. Do not include citations.]',
             'zh-cn': '［请填写 200–300 字摘要，概述本报告的研究目的、方法与主要发现。摘要中不得包含引用。］',
             'zh-tw': '［請填寫 200–300 字摘要，概述本報告的研究目的、方法與主要發現。摘要中不得包含引用。］'}[lang],
        ),
        (
            {'en': 'Keywords', 'zh-cn': '关键词', 'zh-tw': '關鍵詞'}[lang],
            {'en': '5–8 keywords, comma-separated.',
             'zh-cn': '5–8 个关键词，以逗号分隔。',
             'zh-tw': '5–8 個關鍵詞，以逗號分隔。'}[lang],
            {'en': '[keyword 1, keyword 2, keyword 3, keyword 4, keyword 5]',
             'zh-cn': '［关键词一、关键词二、关键词三、关键词四、关键词五］',
             'zh-tw': '［關鍵詞一、關鍵詞二、關鍵詞三、關鍵詞四、關鍵詞五］'}[lang],
        ),
        (
            {'en': 'Table of Contents', 'zh-cn': '目录', 'zh-tw': '目錄'}[lang],
            {'en': 'Auto-generated from headings.',
             'zh-cn': '根据标题自动生成。',
             'zh-tw': '根據標題自動產生。'}[lang],
            {'en': '[Generate automatically from the heading styles used in this document (References → Table of Contents in Word).]',
             'zh-cn': '［请使用本文档中的标题样式自动生成目录（Word 菜单：引用 → 目录）。］',
             'zh-tw': '［請使用本文件中的標題樣式自動產生目錄（Word 選單：參照 → 目錄）。］'}[lang],
        ),
        (
            {'en': 'List of Tables / Figures', 'zh-cn': '图表清单', 'zh-tw': '圖表清單'}[lang],
            {'en': 'Required if the report contains more than three tables or figures.',
             'zh-cn': '若报告包含三个以上图表，则此项为必需。',
             'zh-tw': '若報告包含三個以上圖表，則此項為必需。'}[lang],
            {'en': '[List all tables and figures with captions and page numbers, if applicable.]',
             'zh-cn': '［如适用，请列出全部图表标题及其页码。］',
             'zh-tw': '［如適用，請列出全部圖表標題及其頁碼。］'}[lang],
        ),
        (
            {'en': 'Executive Summary', 'zh-cn': '执行摘要', 'zh-tw': '執行摘要'}[lang],
            {'en': '800–1,200 words; a self-contained summary for non-specialist readers.',
             'zh-cn': '800–1,200 字；面向非专业读者的独立摘要。',
             'zh-tw': '800–1,200 字；面向非專業讀者的獨立摘要。'}[lang],
            {'en': '[Insert an 800–1,200 word executive summary that can be read independently of the full report, written for a non-specialist audience.]',
             'zh-cn': '［请填写 800–1,200 字的执行摘要，内容应可独立于报告正文阅读，并面向非专业读者撰写。］',
             'zh-tw': '［請填寫 800–1,200 字的執行摘要，內容應可獨立於報告正文閱讀，並面向非專業讀者撰寫。］'}[lang],
        ),
        (
            {'en': '1. Introduction', 'zh-cn': '一、引言', 'zh-tw': '一、導言'}[lang],
            {'en': 'Context, rationale, research questions, and report structure.',
             'zh-cn': '研究背景、研究依据、研究问题及报告结构说明。',
             'zh-tw': '研究背景、研究依據、研究問題及報告結構說明。'}[lang],
            {'en': '[Insert the introduction: context, rationale, research questions, and an overview of the report’s structure.]',
             'zh-cn': '［请填写引言：研究背景、研究依据、研究问题，并简要说明报告结构。］',
             'zh-tw': '［請填寫導言：研究背景、研究依據、研究問題，並簡要說明報告結構。］'}[lang],
        ),
        (
            {'en': '2–N. Main Chapters', 'zh-cn': '二至 N、正文章节', 'zh-tw': '二至 N、正文章節'}[lang],
            {'en': 'Methodology, literature, findings, and analysis (structure varies by project).',
             'zh-cn': '研究方法、文献综述、研究发现与分析（具体结构视项目而定）。',
             'zh-tw': '研究方法、文獻綜述、研究發現與分析（具體結構視項目而定）。'}[lang],
            {'en': '[Insert the main body chapters: methodology, literature review, findings, and analysis. Add or remove chapters as appropriate for this project, keeping consistent heading levels.]',
             'zh-cn': '［请填写正文章节，包括研究方法、文献综述、研究发现与分析。可根据本项目需要增减章节，但须保持标题层级一致。］',
             'zh-tw': '［請填寫正文章節，包括研究方法、文獻綜述、研究發現與分析。可根據本項目需要增減章節，但須保持標題層級一致。］'}[lang],
        ),
        (
            {'en': 'Conclusions', 'zh-cn': '结论', 'zh-tw': '結論'}[lang],
            {'en': 'Key findings, policy implications, limitations, and future research.',
             'zh-cn': '主要发现、政策启示、研究局限及未来研究方向。',
             'zh-tw': '主要發現、政策啟示、研究侷限及未來研究方向。'}[lang],
            {'en': '[Summarize the key findings, their policy implications, the study’s limitations, and directions for future research.]',
             'zh-cn': '［请概述主要研究发现、其政策启示、本研究的局限性，以及未来研究方向。］',
             'zh-tw': '［請概述主要研究發現、其政策啟示、本研究的侷限性，以及未來研究方向。］'}[lang],
        ),
        (
            {'en': 'References', 'zh-cn': '参考文献', 'zh-tw': '參考文獻'}[lang],
            {'en': 'APA 7th or Chicago 18th, applied consistently throughout.',
             'zh-cn': '采用 APA 第 7 版或芝加哥格式第 18 版，全文格式须保持一致。',
             'zh-tw': '採用 APA 第 7 版或芝加哥格式第 18 版，全文格式須保持一致。'}[lang],
            {'en': '[Insert the full reference list, formatted consistently in APA 7th or Chicago 18th style.]',
             'zh-cn': '［请填写完整参考文献列表，全文统一采用 APA 第 7 版或芝加哥格式第 18 版。］',
             'zh-tw': '［請填寫完整參考文獻列表，全文統一採用 APA 第 7 版或芝加哥格式第 18 版。］'}[lang],
        ),
        (
            {'en': 'Appendices (if any)', 'zh-cn': '附录（如有）', 'zh-tw': '附錄（如有）'}[lang],
            {'en': 'Data, instruments, or additional tables.',
             'zh-cn': '数据、研究工具或补充图表。',
             'zh-tw': '資料、研究工具或補充圖表。'}[lang],
            {'en': '[Insert any supplementary data, research instruments, or additional tables not included in the main text.]',
             'zh-cn': '［如有，请填写正文之外的补充数据、研究工具或补充图表。］',
             'zh-tw': '［如有，請填寫正文之外的補充資料、研究工具或補充圖表。］'}[lang],
        ),
        (
            {'en': 'About the Authors', 'zh-cn': '作者简介', 'zh-tw': '作者簡介'}[lang],
            {'en': 'Brief biography, 50–100 words per author.',
             'zh-cn': '每位作者简介 50–100 字。',
             'zh-tw': '每位作者簡介 50–100 字。'}[lang],
            {'en': '[Insert a 50–100 word biography for each author, including current affiliation and relevant expertise.]',
             'zh-cn': '［请填写每位作者 50–100 字的简介，包括现任职机构及相关专长。］',
             'zh-tw': '［請填寫每位作者 50–100 字的簡介，包括現任職機構及相關專長。］'}[lang],
        ),
        (
            {'en': 'About the Institute', 'zh-cn': '关于研究院', 'zh-tw': '關於研究院'}[lang],
            L['do_not_alter'][lang],
            L['about_institute_body'][lang],
        ),
    ],
})

# 3. Policy Brief Template ---------------------------------------------------
TEMPLATES.append({
    'slug': 'Policy Brief Template',
    'title': {'en': 'Policy Brief Template', 'zh-cn': '政策简报模板', 'zh-tw': '政策簡報範本'},
    'subtitle': {
        'en': 'Official Template · v1.0 (2026)',
        'zh-cn': '官方模板 · v1.0（2026）',
        'zh-tw': '官方範本 · v1.0（2026）',
    },
    'meta': lambda lang: [
        (L['format'][lang], L['format_v'][lang]),
        (L['version'][lang], L['version_v'][lang]),
        ({'en': 'Typical Length', 'zh-cn': '一般篇幅', 'zh-tw': '一般篇幅'}[lang],
         {'en': '1,500–4,000 words', 'zh-cn': '1,500–4,000 字', 'zh-tw': '1,500–4,000 字'}[lang]),
        ({'en': 'Target Audience', 'zh-cn': '目标受众', 'zh-tw': '目標受眾'}[lang],
         {'en': 'Policymakers, institutional leaders, practitioners', 'zh-cn': '政策制定者、机构领导者、实务工作者', 'zh-tw': '政策制定者、機構領導者、實務工作者'}[lang]),
        (L['submission_lang'][lang], L['submission_lang_v'][lang]),
    ],
    'intro': {
        'en': 'Policy briefs follow a structured format designed to communicate effectively with policy audiences. Unlike academic papers, policy briefs lead with conclusions and emphasize practical recommendations. Write for readers who are knowledgeable but not academic specialists: avoid jargon, define technical terms, and use active sentences. Recommendations should be specific enough that a reader can act on them — for example, "the Ministry of Education should…" rather than "policymakers should consider…".',
        'zh-cn': '政策简报采用旨在有效面向政策受众沟通的结构化格式。与学术论文不同，政策简报以结论开篇，并侧重具有可操作性的建议。撰写时应面向具备一定专业知识但非学术专家的读者：避免使用专业术语，对技术性名词加以说明，并使用主动句式。建议应具体到读者可据此采取行动的程度——例如应写"教育部应……"，而非笼统地写"政策制定者应考虑……"。',
        'zh-tw': '政策簡報採用旨在有效面向政策受眾溝通的結構化格式。與學術論文不同，政策簡報以結論開篇，並側重具有可操作性的建議。撰寫時應面向具備一定專業知識但非學術專家的讀者：避免使用專業術語，對技術性名詞加以說明，並使用主動句式。建議應具體到讀者可據此採取行動的程度——例如應寫「教育部應……」，而非籠統地寫「政策制定者應考慮……」。',
    },
    'sections': lambda lang: [
        (
            {'en': 'Cover', 'zh-cn': '封面', 'zh-tw': '封面'}[lang],
            {'en': 'Title, brief number, author(s), research center, date, DOI.',
             'zh-cn': '标题、简报编号、作者、研究中心、日期、DOI。',
             'zh-tw': '標題、簡報編號、作者、研究中心、日期、DOI。'}[lang],
            {'en': '[Brief Title]\nPolicy Brief No. PB-YYYY-00X\n[Author(s), Affiliation]\n[Research Center]\n[Date]\n[DOI: 10.XXXXX/YYYYY]',
             'zh-cn': '［简报标题］\n政策简报编号 PB-YYYY-00X\n［作者、所属机构］\n［研究中心］\n［日期］\n［DOI：10.XXXXX/YYYYY］',
             'zh-tw': '［簡報標題］\n政策簡報編號 PB-YYYY-00X\n［作者、所屬機構］\n［研究中心］\n［日期］\n［DOI：10.XXXXX/YYYYY］'}[lang],
        ),
        (
            {'en': 'Key Messages', 'zh-cn': '核心信息', 'zh-tw': '核心資訊'}[lang],
            {'en': '3–5 bullet points; the brief’s essential conclusions, appearing before the text body.',
             'zh-cn': '3–5 条要点，呈现简报的核心结论，置于正文之前。',
             'zh-tw': '3–5 條要點，呈現簡報的核心結論，置於正文之前。'}[lang],
            {'en': '[• Key message 1]\n[• Key message 2]\n[• Key message 3]',
             'zh-cn': '［• 核心信息一］\n［• 核心信息二］\n［• 核心信息三］',
             'zh-tw': '［• 核心資訊一］\n［• 核心資訊二］\n［• 核心資訊三］'}[lang],
        ),
        (
            {'en': 'Executive Summary', 'zh-cn': '执行摘要', 'zh-tw': '執行摘要'}[lang],
            {'en': '150–250 words, for readers who will read no further.',
             'zh-cn': '150–250 字，供不再继续阅读全文的读者参考。',
             'zh-tw': '150–250 字，供不再繼續閱讀全文的讀者參考。'}[lang],
            {'en': '[Insert a 150–250 word summary that stands alone for readers who read no further than this section.]',
             'zh-cn': '［请填写 150–250 字摘要，内容应可独立成篇，供仅阅读此部分的读者理解全貌。］',
             'zh-tw': '［請填寫 150–250 字摘要，內容應可獨立成篇，供僅閱讀此部分的讀者理解全貌。］'}[lang],
        ),
        (
            {'en': '1. Policy Context', 'zh-cn': '一、政策背景', 'zh-tw': '一、政策背景'}[lang],
            {'en': 'The policy problem or question the brief addresses, and why it matters now.',
             'zh-cn': '本简报所针对的政策问题及其当前重要性。',
             'zh-tw': '本簡報所針對的政策問題及其當前重要性。'}[lang],
            {'en': '[Describe the policy problem or question this brief addresses, and explain why it matters now.]',
             'zh-cn': '［请描述本简报所针对的政策问题，并说明其在当下的重要性。］',
             'zh-tw': '［請描述本簡報所針對的政策問題，並說明其在當下的重要性。］'}[lang],
        ),
        (
            {'en': '2. Evidence and Analysis', 'zh-cn': '二、证据与分析', 'zh-tw': '二、證據與分析'}[lang],
            {'en': 'The evidence base and structured analysis, clearly distinguished from opinion.',
             'zh-cn': '证据基础与结构化分析，须与个人观点清楚区分。',
             'zh-tw': '證據基礎與結構化分析，須與個人觀點清楚區分。'}[lang],
            {'en': '[Present the evidence base and analysis supporting this brief. Clearly distinguish evidence from opinion or interpretation.]',
             'zh-cn': '［请呈现支撑本简报的证据基础与分析内容，并将证据与观点或诠释清楚区分。］',
             'zh-tw': '［請呈現支撐本簡報的證據基礎與分析內容，並將證據與觀點或詮釋清楚區分。］'}[lang],
        ),
        (
            {'en': '3. Policy Options', 'zh-cn': '三、政策选项', 'zh-tw': '三、政策選項'}[lang],
            {'en': 'Realistic options for action, with the advantages and limitations of each.',
             'zh-cn': '切实可行的行动选项，并说明各选项的优势与局限。',
             'zh-tw': '切實可行的行動選項，並說明各選項的優勢與侷限。'}[lang],
            {'en': '[Present realistic policy options, with the advantages and limitations of each.]',
             'zh-cn': '［请呈现切实可行的政策选项，并说明各选项的优势与局限。］',
             'zh-tw': '［請呈現切實可行的政策選項，並說明各選項的優勢與侷限。］'}[lang],
        ),
        (
            {'en': '4. Recommendations', 'zh-cn': '四、建议', 'zh-tw': '四、建議'}[lang],
            {'en': 'Specific, actionable recommendations addressed to a named audience.',
             'zh-cn': '具体、可操作的建议，并指明适用对象。',
             'zh-tw': '具體、可操作的建議，並指明適用對象。'}[lang],
            {'en': '[Insert specific, actionable recommendations, each addressed to a named audience (e.g., a specific ministry, agency, or institution).]',
             'zh-cn': '［请填写具体且可操作的建议，并逐条注明适用对象（如具体的部委、机构或组织）。］',
             'zh-tw': '［請填寫具體且可操作的建議，並逐條註明適用對象（如具體的部會、機構或組織）。］'}[lang],
        ),
        (
            {'en': 'References', 'zh-cn': '参考文献', 'zh-tw': '參考文獻'}[lang],
            {'en': 'Essential references only, inline or endnote style — not a full bibliography.',
             'zh-cn': '仅列出核心参考文献，采用文中夹注或尾注格式，无需完整参考书目。',
             'zh-tw': '僅列出核心參考文獻，採用文中夾註或尾註格式，無需完整參考書目。'}[lang],
            {'en': '[List only the essential references cited in this brief, using inline or endnote style.]',
             'zh-cn': '［请仅列出本简报引用的核心参考文献，采用文中夹注或尾注格式。］',
             'zh-tw': '［請僅列出本簡報引用的核心參考文獻，採用文中夾註或尾註格式。］'}[lang],
        ),
        (
            {'en': 'About the Institute', 'zh-cn': '关于研究院', 'zh-tw': '關於研究院'}[lang],
            L['do_not_alter'][lang],
            L['about_institute_body'][lang],
        ),
    ],
})

# 4. Working Paper Template ---------------------------------------------------
TEMPLATES.append({
    'slug': 'Working Paper Template',
    'title': {'en': 'Working Paper Template', 'zh-cn': '工作论文模板', 'zh-tw': '工作論文範本'},
    'subtitle': {
        'en': 'Official Template · v1.0 (2026)',
        'zh-cn': '官方模板 · v1.0（2026）',
        'zh-tw': '官方範本 · v1.0（2026）',
    },
    'meta': lambda lang: [
        (L['format'][lang], L['format_v'][lang]),
        (L['version'][lang], L['version_v'][lang]),
        ({'en': 'Typical Length', 'zh-cn': '一般篇幅', 'zh-tw': '一般篇幅'}[lang],
         {'en': '5,000–15,000 words', 'zh-cn': '5,000–15,000 字', 'zh-tw': '5,000–15,000 字'}[lang]),
        ({'en': 'DOI', 'zh-cn': 'DOI（数字对象标识符）', 'zh-tw': 'DOI（數位物件識別碼）'}[lang],
         {'en': 'Registered (Crossref)', 'zh-cn': '已注册（Crossref）', 'zh-tw': '已註冊（Crossref）'}[lang]),
        (L['submission_lang'][lang], L['submission_lang_v'][lang]),
    ],
    'intro': {
        'en': 'Working papers enable rapid dissemination of research findings, theoretical contributions, and scholarly arguments before or alongside the Institute’s formal peer review process. Working papers may present preliminary findings, theoretical frameworks, systematic literature reviews, or commentary on emerging developments in the Institute’s research areas. See the Citation Guide for the PSG Author-Date format used in Institute references (Working Paper No. WP-YYYY-00X).',
        'zh-cn': '工作论文用于在研究院正式同行评审流程之前或同时，快速传播研究发现、理论贡献与学术论点。工作论文可呈现初步研究发现、理论框架、系统性文献综述，或对研究院相关研究领域最新进展的评论。研究院参考文献采用 PSG 作者-出版年格式，编号格式为"工作论文 No. WP-YYYY-00X"，具体规则请参阅引用指南。',
        'zh-tw': '工作論文用於在研究院正式同儕評審流程之前或同時，快速傳播研究發現、理論貢獻與學術論點。工作論文可呈現初步研究發現、理論框架、系統性文獻綜述，或對研究院相關研究領域最新進展的評論。研究院參考文獻採用 PSG 作者-出版年格式，編號格式為「工作論文 No. WP-YYYY-00X」，具體規則請參閱引用指南。',
    },
    'sections': lambda lang: [
        (
            {'en': 'Cover Page', 'zh-cn': '封面', 'zh-tw': '封面'}[lang],
            {'en': 'Title, working paper number, author(s) and affiliation, date, DOI.',
             'zh-cn': '标题、工作论文编号、作者及所属机构、日期、DOI。',
             'zh-tw': '標題、工作論文編號、作者及所屬機構、日期、DOI。'}[lang],
            {'en': '[Paper Title]\nWorking Paper No. WP-YYYY-00X\n[Author(s), Affiliation]\n[Date]\n[DOI: 10.XXXXX/YYYYY]',
             'zh-cn': '［论文标题］\n工作论文编号 WP-YYYY-00X\n［作者、所属机构］\n［日期］\n［DOI：10.XXXXX/YYYYY］',
             'zh-tw': '［論文標題］\n工作論文編號 WP-YYYY-00X\n［作者、所屬機構］\n［日期］\n［DOI：10.XXXXX/YYYYY］'}[lang],
        ),
        (
            {'en': 'Abstract', 'zh-cn': '摘要', 'zh-tw': '摘要'}[lang],
            {'en': '150–250 words; no citations in the abstract.',
             'zh-cn': '150–250 字；摘要中不得包含引用。',
             'zh-tw': '150–250 字；摘要中不得包含引用。'}[lang],
            {'en': '[Insert a 150–250 word abstract summarizing the paper’s argument, methods (if applicable), and contribution. Do not include citations.]',
             'zh-cn': '［请填写 150–250 字摘要，概述本文的论点、研究方法（如适用）及学术贡献。摘要中不得包含引用。］',
             'zh-tw': '［請填寫 150–250 字摘要，概述本文的論點、研究方法（如適用）及學術貢獻。摘要中不得包含引用。］'}[lang],
        ),
        (
            {'en': 'Keywords', 'zh-cn': '关键词', 'zh-tw': '關鍵詞'}[lang],
            {'en': '4–6 keywords, comma-separated.',
             'zh-cn': '4–6 个关键词，以逗号分隔。',
             'zh-tw': '4–6 個關鍵詞，以逗號分隔。'}[lang],
            {'en': '[keyword 1, keyword 2, keyword 3, keyword 4]',
             'zh-cn': '［关键词一、关键词二、关键词三、关键词四］',
             'zh-tw': '［關鍵詞一、關鍵詞二、關鍵詞三、關鍵詞四］'}[lang],
        ),
        (
            {'en': '1. Introduction', 'zh-cn': '一、引言', 'zh-tw': '一、導言'}[lang],
            {'en': 'The research problem, motivation, and the paper’s contribution.',
             'zh-cn': '研究问题、研究动机及本文的学术贡献。',
             'zh-tw': '研究問題、研究動機及本文的學術貢獻。'}[lang],
            {'en': '[Insert the introduction: the research problem, motivation, and what this paper contributes to the field.]',
             'zh-cn': '［请填写引言：研究问题、研究动机，以及本文对该领域的学术贡献。］',
             'zh-tw': '［請填寫導言：研究問題、研究動機，以及本文對該領域的學術貢獻。］'}[lang],
        ),
        (
            {'en': '2. Literature Review / Theoretical Framework', 'zh-cn': '二、文献综述／理论框架', 'zh-tw': '二、文獻綜述／理論框架'}[lang],
            {'en': 'Positioning the paper within existing scholarship.',
             'zh-cn': '将本文置于既有学术研究脉络中加以定位。',
             'zh-tw': '將本文置於既有學術研究脈絡中加以定位。'}[lang],
            {'en': '[Review the relevant literature or set out the theoretical framework, positioning this paper within the existing scholarship.]',
             'zh-cn': '［请综述相关文献或阐述理论框架，将本文置于既有学术研究脉络中加以定位。］',
             'zh-tw': '［請綜述相關文獻或闡述理論框架，將本文置於既有學術研究脈絡中加以定位。］'}[lang],
        ),
        (
            {'en': '3. Methodology', 'zh-cn': '三、研究方法', 'zh-tw': '三、研究方法'}[lang],
            {'en': 'Research design and methods, where applicable.',
             'zh-cn': '研究设计与方法（如适用）。',
             'zh-tw': '研究設計與方法（如適用）。'}[lang],
            {'en': '[Describe the research design and methods used, if this paper reports original empirical research. Omit this section for purely theoretical or conceptual papers.]',
             'zh-cn': '［如本文包含原创实证研究，请描述所采用的研究设计与方法；纯理论或概念性文章可省略本节。］',
             'zh-tw': '［如本文包含原創實證研究，請描述所採用的研究設計與方法；純理論或概念性文章可省略本節。］'}[lang],
        ),
        (
            {'en': '4. Findings / Argument', 'zh-cn': '四、研究发现／论点', 'zh-tw': '四、研究發現／論點'}[lang],
            {'en': 'The main analysis, evidence, or theoretical argument.',
             'zh-cn': '主要分析内容、证据或理论论点。',
             'zh-tw': '主要分析內容、證據或理論論點。'}[lang],
            {'en': '[Present the paper’s main analysis, evidence, or theoretical argument. Structure with sub-headings as needed.]',
             'zh-cn': '［请呈现本文的主要分析、证据或理论论点，可视需要使用子标题划分层次。］',
             'zh-tw': '［請呈現本文的主要分析、證據或理論論點，可視需要使用子標題劃分層次。］'}[lang],
        ),
        (
            {'en': '5. Discussion', 'zh-cn': '五、讨论', 'zh-tw': '五、討論'}[lang],
            {'en': 'Interpretation, implications, and limitations.',
             'zh-cn': '结果诠释、意义与研究局限。',
             'zh-tw': '結果詮釋、意義與研究侷限。'}[lang],
            {'en': '[Interpret the findings or argument, discuss their implications, and note the paper’s limitations.]',
             'zh-cn': '［请诠释研究发现或论点，讨论其意义，并说明本文的研究局限。］',
             'zh-tw': '［請詮釋研究發現或論點，討論其意義，並說明本文的研究侷限。］'}[lang],
        ),
        (
            {'en': 'Conclusion', 'zh-cn': '结论', 'zh-tw': '結論'}[lang],
            {'en': 'Summary of the contribution and directions for further work.',
             'zh-cn': '总结学术贡献并指出后续研究方向。',
             'zh-tw': '總結學術貢獻並指出後續研究方向。'}[lang],
            {'en': '[Summarize this paper’s contribution and suggest directions for further research.]',
             'zh-cn': '［请总结本文的学术贡献，并提出后续研究方向的建议。］',
             'zh-tw': '［請總結本文的學術貢獻，並提出後續研究方向的建議。］'}[lang],
        ),
        (
            {'en': 'References', 'zh-cn': '参考文献', 'zh-tw': '參考文獻'}[lang],
            {'en': 'PSG Author-Date format (see the Citation Guide).',
             'zh-cn': '采用 PSG 作者-出版年格式（详见引用指南）。',
             'zh-tw': '採用 PSG 作者-出版年格式（詳見引用指南）。'}[lang],
            {'en': '[Insert the full reference list in PSG Author-Date format.]',
             'zh-cn': '［请填写完整参考文献列表，采用 PSG 作者-出版年格式。］',
             'zh-tw': '［請填寫完整參考文獻列表，採用 PSG 作者-出版年格式。］'}[lang],
        ),
        (
            {'en': 'About the Author(s)', 'zh-cn': '作者简介', 'zh-tw': '作者簡介'}[lang],
            {'en': 'Brief biography, 50–100 words per author.',
             'zh-cn': '每位作者简介 50–100 字。',
             'zh-tw': '每位作者簡介 50–100 字。'}[lang],
            {'en': '[Insert a 50–100 word biography for each author, including current affiliation and relevant expertise.]',
             'zh-cn': '［请填写每位作者 50–100 字的简介，包括现任职机构及相关专长。］',
             'zh-tw': '［請填寫每位作者 50–100 字的簡介，包括現任職機構及相關專長。］'}[lang],
        ),
        (
            {'en': 'About the Institute', 'zh-cn': '关于研究院', 'zh-tw': '關於研究院'}[lang],
            L['do_not_alter'][lang],
            L['about_institute_body'][lang],
        ),
    ],
})

# 5. Event Minutes Summary Template -------------------------------------------
TEMPLATES.append({
    'slug': 'Event Proceedings Summary Template',
    'title': {'en': 'Event Proceedings Summary Template', 'zh-cn': '活动会议记录摘要模板', 'zh-tw': '活動會議記錄摘要範本'},
    'subtitle': {
        'en': 'Official Template · v1.0 (2026)',
        'zh-cn': '官方模板 · v1.0（2026）',
        'zh-tw': '官方範本 · v1.0（2026）',
    },
    'meta': lambda lang: [
        (L['format'][lang], L['format_v'][lang]),
        (L['version'][lang], L['version_v'][lang]),
        ({'en': 'Typical Length', 'zh-cn': '一般篇幅', 'zh-tw': '一般篇幅'}[lang],
         {'en': '1,000–3,000 words (varies by event)', 'zh-cn': '1,000–3,000 字（依活动规模而定）', 'zh-tw': '1,000–3,000 字（依活動規模而定）'}[lang]),
        ({'en': 'Prepared By', 'zh-cn': '编写人', 'zh-tw': '編寫人'}[lang],
         {'en': 'Event organizing team or designated rapporteur', 'zh-cn': '活动筹办团队或指定记录人', 'zh-tw': '活動籌辦團隊或指定記錄人'}[lang]),
        (L['submission_lang'][lang], L['submission_lang_v'][lang]),
    ],
    'intro': {
        'en': 'This template is used to prepare a summary record of an Institute conference, seminar, workshop, or training program after it concludes. Completed summaries support the Institute’s events archive and may inform future publications or programming. Submit the completed summary to the organizing research center within two weeks of the event.',
        'zh-cn': '本模板用于在研究院会议、研讨会、工作坊或培训项目结束后，编制活动会议记录摘要。完成后的摘要将纳入研究院活动档案，并可为后续出版物或活动策划提供参考。请在活动结束后两周内，将完成的摘要提交至主办研究中心。',
        'zh-tw': '本範本用於在研究院會議、研討會、工作坊或培訓項目結束後，編製活動會議記錄摘要。完成後的摘要將納入研究院活動檔案，並可為後續出版物或活動策劃提供參考。請於活動結束後兩週內，將完成的摘要提交至主辦研究中心。',
    },
    'sections': lambda lang: [
        (
            {'en': 'Cover', 'zh-cn': '封面', 'zh-tw': '封面'}[lang],
            {'en': 'Event title, date, format, organizing research center, prepared by, date of minutes.',
             'zh-cn': '活动名称、日期、举办形式、主办研究中心、编写人、记录日期。',
             'zh-tw': '活動名稱、日期、舉辦形式、主辦研究中心、編寫人、記錄日期。'}[lang],
            {'en': '[Event Title]\n[Date]\n[Format: Online / In-Person / Hybrid]\n[Organizing Research Center]\n[Prepared By]\n[Date of Minutes]',
             'zh-cn': '［活动名称］\n［日期］\n［举办形式：线上／线下／混合］\n［主办研究中心］\n［编写人］\n［记录日期］',
             'zh-tw': '［活動名稱］\n［日期］\n［舉辦形式：線上／線下／混合］\n［主辦研究中心］\n［編寫人］\n［記錄日期］'}[lang],
        ),
        (
            {'en': '1. Event Overview', 'zh-cn': '一、活动概况', 'zh-tw': '一、活動概況'}[lang],
            {'en': 'Purpose, format, and attendance summary.',
             'zh-cn': '活动目的、举办形式及参会情况概述。',
             'zh-tw': '活動目的、舉辦形式及與會情況概述。'}[lang],
            {'en': '[Describe the purpose of the event, its format, and a brief summary of attendance (approximate number and composition of participants).]',
             'zh-cn': '［请描述本次活动的目的、举办形式，并简要概述参会情况（参会人数及构成）。］',
             'zh-tw': '［請描述本次活動的目的、舉辦形式，並簡要概述與會情況（與會人數及構成）。］'}[lang],
        ),
        (
            {'en': '2. Speaker and Session Summaries', 'zh-cn': '二、演讲者与场次摘要', 'zh-tw': '二、講者與場次摘要'}[lang],
            {'en': 'Repeat this block for each session: speaker name and affiliation, session title, and a 150–250 word summary of key points.',
             'zh-cn': '每个场次重复填写本区块：演讲者姓名及所属机构、场次标题，以及 150–250 字的要点摘要。',
             'zh-tw': '每個場次重複填寫本區塊：講者姓名及所屬機構、場次標題，以及 150–250 字的要點摘要。'}[lang],
            {'en': '[Session Title]\n[Speaker Name, Affiliation]\n[150–250 word summary of the session’s key points]\n\n(Repeat this block for each additional session.)',
             'zh-cn': '［场次标题］\n［演讲者姓名、所属机构］\n［150–250 字的场次要点摘要］\n\n（如有多个场次，请重复填写本区块。）',
             'zh-tw': '［場次標題］\n［講者姓名、所屬機構］\n［150–250 字的場次要點摘要］\n\n（如有多個場次，請重複填寫本區塊。）'}[lang],
        ),
        (
            {'en': '3. Discussion Highlights', 'zh-cn': '三、讨论要点', 'zh-tw': '三、討論要點'}[lang],
            {'en': 'Key questions raised and responses given during Q&A or discussion.',
             'zh-cn': '问答或讨论环节中提出的主要问题及回应。',
             'zh-tw': '問答或討論環節中提出的主要問題及回應。'}[lang],
            {'en': '[Summarize the key questions raised and the responses given during any Q&A or open discussion segments.]',
             'zh-cn': '［请概述问答或自由讨论环节中提出的主要问题及相应回应。］',
             'zh-tw': '［請概述問答或自由討論環節中提出的主要問題及相應回應。］'}[lang],
        ),
        (
            {'en': '4. Key Conclusions and Outcomes', 'zh-cn': '四、主要结论与成果', 'zh-tw': '四、主要結論與成果'}[lang],
            {'en': 'Main takeaways and any decisions reached.',
             'zh-cn': '主要收获及达成的决定（如有）。',
             'zh-tw': '主要收穫及達成的決定（如有）。'}[lang],
            {'en': '[Summarize the main takeaways from the event and any decisions or agreements reached.]',
             'zh-cn': '［请概述本次活动的主要收获，以及达成的任何决定或共识。］',
             'zh-tw': '［請概述本次活動的主要收穫，以及達成的任何決定或共識。］'}[lang],
        ),
        (
            {'en': '5. Follow-up Actions', 'zh-cn': '五、后续行动', 'zh-tw': '五、後續行動'}[lang],
            {'en': 'Action items, owners, and target dates, if any.',
             'zh-cn': '后续行动事项、负责人及目标完成日期（如有）。',
             'zh-tw': '後續行動事項、負責人及目標完成日期（如有）。'}[lang],
            {'en': '[List any follow-up action items, the person or team responsible, and the target completion date. State "No follow-up actions identified" if none.]',
             'zh-cn': '［如有，请列出后续行动事项、负责人／团队及目标完成日期；若无后续行动，请填写"无后续行动事项"。］',
             'zh-tw': '［如有，請列出後續行動事項、負責人／團隊及目標完成日期；若無後續行動，請填寫「無後續行動事項」。］'}[lang],
        ),
        (
            {'en': 'Appendix: Programme and Attendance', 'zh-cn': '附录：议程与出席名单', 'zh-tw': '附錄：議程與出席名單'}[lang],
            {'en': 'Full programme schedule and attendee list (optional).',
             'zh-cn': '完整活动议程及出席人员名单（可选）。',
             'zh-tw': '完整活動議程及出席人員名單（可選）。'}[lang],
            {'en': '[Attach the full programme schedule and, if appropriate, a list of attendees. This section is optional.]',
             'zh-cn': '［如适用，请附上完整活动议程及出席人员名单。本节为可选内容。］',
             'zh-tw': '［如適用，請附上完整活動議程及出席人員名單。本節為可選內容。］'}[lang],
        ),
    ],
})

# 6. Edited Volume Chapter Template -------------------------------------------
TEMPLATES.append({
    'slug': 'Edited Volume Chapter Template',
    'title': {'en': 'Edited Volume Chapter Template', 'zh-cn': '主编文集章节模板', 'zh-tw': '主編文集章節範本'},
    'subtitle': {
        'en': 'Official Template · v1.0 (2026)',
        'zh-cn': '官方模板 · v1.0（2026）',
        'zh-tw': '官方範本 · v1.0（2026）',
    },
    'meta': lambda lang: [
        (L['format'][lang], L['format_v'][lang]),
        (L['version'][lang], L['version_v'][lang]),
        ({'en': 'Typical Length', 'zh-cn': '一般篇幅', 'zh-tw': '一般篇幅'}[lang],
         {'en': '6,000–9,000 words (per the call for chapters)', 'zh-cn': '6,000–9,000 字（以具体征稿启事为准）', 'zh-tw': '6,000–9,000 字（以具體徵稿啟事為準）'}[lang]),
        ({'en': 'Citation Style', 'zh-cn': '引用格式', 'zh-tw': '引用格式'}[lang],
         {'en': 'PSG Author-Date Format', 'zh-cn': 'PSG 作者-出版年格式', 'zh-tw': 'PSG 作者-出版年格式'}[lang]),
        (L['submission_lang'][lang], L['submission_lang_v'][lang]),
    ],
    'intro': {
        'en': 'This template is used to prepare a chapter for submission to an edited volume published in association with Panorama Research Institute, once a proposed chapter has been accepted following a call for chapters. Follow the specific volume’s editorial guidelines for exact word limits and formatting requirements; this template reflects the Institute’s general house style.',
        'zh-cn': '本模板用于在章节提案通过研究院相关主编文集征稿启事评审后，撰写并提交正式章节。具体字数要求及格式规范请以该文集的编辑指南为准；本模板体现研究院的通用格式规范。',
        'zh-tw': '本範本用於在章節提案通過研究院相關主編文集徵稿啟事評審後，撰寫並提交正式章節。具體字數要求及格式規範請以該文集的編輯指南為準；本範本體現研究院的通用格式規範。',
    },
    'sections': lambda lang: [
        (
            {'en': 'Cover Page', 'zh-cn': '封面', 'zh-tw': '封面'}[lang],
            {'en': 'Chapter title, volume title, author(s) and affiliation, word count.',
             'zh-cn': '章节标题、文集标题、作者及所属机构、字数统计。',
             'zh-tw': '章節標題、文集標題、作者及所屬機構、字數統計。'}[lang],
            {'en': '[Chapter Title]\nIn: [Volume Title]\n[Author(s), Affiliation]\n[Word Count]',
             'zh-cn': '［章节标题］\n收录于：［文集标题］\n［作者、所属机构］\n［字数］',
             'zh-tw': '［章節標題］\n收錄於：［文集標題］\n［作者、所屬機構］\n［字數］'}[lang],
        ),
        (
            {'en': 'Abstract', 'zh-cn': '摘要', 'zh-tw': '摘要'}[lang],
            {'en': '150–200 words.',
             'zh-cn': '150–200 字。',
             'zh-tw': '150–200 字。'}[lang],
            {'en': '[Insert a 150–200 word abstract summarizing the chapter’s argument and contribution to the volume.]',
             'zh-cn': '［请填写 150–200 字摘要，概述本章节的论点及其对文集的贡献。］',
             'zh-tw': '［請填寫 150–200 字摘要，概述本章節的論點及其對文集的貢獻。］'}[lang],
        ),
        (
            {'en': 'Keywords', 'zh-cn': '关键词', 'zh-tw': '關鍵詞'}[lang],
            {'en': '4–6 keywords, comma-separated.',
             'zh-cn': '4–6 个关键词，以逗号分隔。',
             'zh-tw': '4–6 個關鍵詞，以逗號分隔。'}[lang],
            {'en': '[keyword 1, keyword 2, keyword 3, keyword 4]',
             'zh-cn': '［关键词一、关键词二、关键词三、关键词四］',
             'zh-tw': '［關鍵詞一、關鍵詞二、關鍵詞三、關鍵詞四］'}[lang],
        ),
        (
            {'en': '1. Introduction', 'zh-cn': '一、引言', 'zh-tw': '一、導言'}[lang],
            {'en': 'Positioning within the volume’s theme and scope.',
             'zh-cn': '将本章节置于文集主题与范围之中加以定位。',
             'zh-tw': '將本章節置於文集主題與範圍之中加以定位。'}[lang],
            {'en': '[Introduce the chapter’s topic and explain how it relates to the volume’s overall theme and scope.]',
             'zh-cn': '［请介绍本章节主题，并说明其与文集整体主题及范围的关联。］',
             'zh-tw': '［請介紹本章節主題，並說明其與文集整體主題及範圍的關聯。］'}[lang],
        ),
        (
            {'en': '2–N. Main Sections', 'zh-cn': '二至 N、正文各节', 'zh-tw': '二至 N、正文各節'}[lang],
            {'en': 'Argument, evidence, and analysis; section structure is set by the author within the volume’s guidelines.',
             'zh-cn': '论点、证据与分析；具体节次结构由作者在文集编辑指南范围内自行设定。',
             'zh-tw': '論點、證據與分析；具體節次結構由作者在文集編輯指南範圍內自行設定。'}[lang],
            {'en': '[Insert the main body of the chapter: argument, evidence, and analysis, divided into sections as appropriate for this chapter’s content.]',
             'zh-cn': '［请填写章节正文：论点、证据与分析，可根据本章节内容自行划分节次。］',
             'zh-tw': '［請填寫章節正文：論點、證據與分析，可根據本章節內容自行劃分節次。］'}[lang],
        ),
        (
            {'en': 'Conclusion', 'zh-cn': '结论', 'zh-tw': '結論'}[lang],
            {'en': 'Summary of the chapter’s contribution to the volume.',
             'zh-cn': '总结本章节对文集的贡献。',
             'zh-tw': '總結本章節對文集的貢獻。'}[lang],
            {'en': '[Summarize this chapter’s main contribution and its significance to the volume as a whole.]',
             'zh-cn': '［请总结本章节的主要贡献及其对文集整体的意义。］',
             'zh-tw': '［請總結本章節的主要貢獻及其對文集整體的意義。］'}[lang],
        ),
        (
            {'en': 'References', 'zh-cn': '参考文献', 'zh-tw': '參考文獻'}[lang],
            {'en': 'PSG Author-Date format (see the Citation Guide).',
             'zh-cn': '采用 PSG 作者-出版年格式（详见引用指南）。',
             'zh-tw': '採用 PSG 作者-出版年格式（詳見引用指南）。'}[lang],
            {'en': '[Insert the full reference list in PSG Author-Date format.]',
             'zh-cn': '［请填写完整参考文献列表，采用 PSG 作者-出版年格式。］',
             'zh-tw': '［請填寫完整參考文獻列表，採用 PSG 作者-出版年格式。］'}[lang],
        ),
        (
            {'en': 'Author Biography', 'zh-cn': '作者简介', 'zh-tw': '作者簡介'}[lang],
            {'en': '75–100 words per author.',
             'zh-cn': '每位作者简介 75–100 字。',
             'zh-tw': '每位作者簡介 75–100 字。'}[lang],
            {'en': '[Insert a 75–100 word biography for each author, including current affiliation and relevant expertise.]',
             'zh-cn': '［请填写每位作者 75–100 字的简介，包括现任职机构及相关专长。］',
             'zh-tw': '［請填寫每位作者 75–100 字的簡介，包括現任職機構及相關專長。］'}[lang],
        ),
        (
            {'en': 'Editorial Note', 'zh-cn': '编辑说明', 'zh-tw': '編輯說明'}[lang],
            None,
            {'en': 'Typical chapter length is 6,000–9,000 words unless the call for chapters states otherwise. Submit the completed chapter to the volume’s editor by the deadline stated in the call for chapters.',
             'zh-cn': '除征稿启事另有说明外，章节篇幅一般为 6,000–9,000 字。请在征稿启事所列截止日期前，将完成的章节提交至该文集主编。',
             'zh-tw': '除徵稿啟事另有說明外，章節篇幅一般為 6,000–9,000 字。請於徵稿啟事所列截止日期前，將完成的章節提交至該文集主編。'}[lang],
        ),
    ],
})


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    if not os.path.exists(LOGO_HEADER) or not os.path.exists(LOGO_COVER):
        lines = [
            'Logo assets not found. Expected:',
            '  ' + LOGO_HEADER,
            '  ' + LOGO_COVER,
            '',
            'Regenerate them by rasterizing public/brand/logo-mono-black.svg (for the',
            'header) and public/brand/logo-stacked.svg (for the cover) to PNG with',
            "Node's `sharp` package, trim the transparent padding, and save them as",
            'logo-header.png / logo-cover.png in a directory of your choosing. Then',
            'either set the PRI_LOGO_ASSETS_DIR environment variable to that directory',
            'before running this script, or place the two PNGs in .logo-assets/ at the',
            'repo root (the default).',
        ]
        raise SystemExit('\n'.join(lines))

    count = 0
    for tpl in TEMPLATES:
        for lang in LANGS:
            title = tpl['title'][lang]
            subtitle = tpl['subtitle'][lang]
            meta_rows = tpl['meta'](lang)
            intro = tpl['intro'][lang]
            sections = tpl['sections'](lang)
            doc = build_doc(lang, title, subtitle, meta_rows, intro, sections)
            filename = f"PRI_{tpl['slug']}_{LANG_SUFFIX[lang]}.docx"
            out_path = os.path.join(OUT_DIR, filename)
            doc.save(out_path)
            print(f'Wrote {filename}')
            count += 1

    print(f'\nGenerated {count} template documents in {OUT_DIR}')


if __name__ == '__main__':
    main()
